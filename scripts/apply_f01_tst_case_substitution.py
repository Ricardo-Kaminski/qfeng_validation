"""Apply F0-1/C-6 fix: substitute fabricated TST case with verified TST-Ag-RR-868-65.2021.5.13.0030."""
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOCX_IN = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")
DOCX_OUT = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")

REPLACEMENTS = [
    # Case number variants
    ("TST-RR-000200-50.2019.5.02.0020",      "TST-Ag-RR-868-65.2021.5.13.0030"),
    ("RR-000200-50.2019.5.02.0020",           "Ag-RR-868-65.2021.5.13.0030"),
    ("RR-0000200-50.2019.5.02.0020",          "Ag-RR-868-65.2021.5.13.0030"),
    # File reference
    ("tst_rr_000200_50_2019.lp",              "tst_ag_rr_868_65_2021.lp"),
    ("tst_rr_000200_50_2019",                 "tst_ag_rr_868_65_2021"),
    # Predicate string used in facts file
    ("TST_RR_000200_50_2019_5_02_0020",       "TST_Ag_RR_868_65_2021_5_13_0030"),
    # Year reference in context of this case
    ("decisao_tst_fundamentada_2019",         "decisao_tst_fundamentada_2023"),
    # Add Tema 1046 context where the case is cited without it
    (
        "TST-Ag-RR-868-65.2021.5.13.0030 (fundamentação completa)",
        "TST-Ag-RR-868-65.2021.5.13.0030 (2ª Turma, DEJT 06/12/2023 — Tema 1046/STF)",
    ),
]


def replace_in_paragraph(para, old: str, new: str) -> int:
    """Replace old→new in a paragraph's runs, handling cross-run splits."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return 0
    replaced = full.replace(old, new)
    # Rewrite: put all text into first run, blank the rest
    if para.runs:
        para.runs[0].text = replaced
        for r in para.runs[1:]:
            r.text = ""
    return full.count(old)


def _iter_all_paragraphs(doc: Document):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def apply_replacements(doc: Document, dry_run: bool = False):
    total = 0

    for old, new in REPLACEMENTS:
        count_run = 0
        for para in _iter_all_paragraphs(doc):
            joined = "".join(r.text for r in para.runs)
            if old in joined:
                if not dry_run:
                    count_run += replace_in_paragraph(para, old, new)
                else:
                    count_run += joined.count(old)

        # XML-level fallback for split-run occurrences
        count_xml = 0
        raw_xml = doc._element.xml
        if old in raw_xml and not dry_run:
            # Replace directly in the lxml element tree text nodes
            from lxml import etree
            for node in doc._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if old in (node.text or ""):
                    node.text = node.text.replace(old, new)
                    count_xml += 1
        elif old in raw_xml and dry_run:
            count_xml = raw_xml.count(f">{old}<") + raw_xml.count(f">{old} ")

        count = count_run + count_xml
        if count:
            label = "[DRY]" if dry_run else "[OK]"
            suffix = f" (xml:{count_xml})" if count_xml else ""
            print(f"  {label} '{old[:50]}' → '{new[:50]}' ({count}×{suffix})")
            total += count
        elif old in raw_xml:
            print(f"  [WARN] '{old[:50]}' in raw XML but extraction failed — check manually")
    return total


def main(dry_run: bool = False):
    if not DOCX_IN.exists():
        raise FileNotFoundError(f"Input not found: {DOCX_IN}")

    bak = DOCX_IN.with_suffix(f".bak_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    if not dry_run:
        shutil.copy2(DOCX_IN, bak)
        print(f"Backup: {bak}")

    doc = Document(DOCX_IN)
    print(f"\n{'DRY RUN — ' if dry_run else ''}Applying F0-1/C-6 TST case substitution...")
    total = apply_replacements(doc, dry_run=dry_run)

    if not dry_run:
        doc.save(DOCX_OUT)
        print(f"\nSaved: {DOCX_OUT}")

    print(f"\nTotal replacements: {total}")

    # Post-check
    doc2 = Document(DOCX_OUT if not dry_run else DOCX_IN)
    all_text = " ".join("".join(r.text for r in p.runs) for p in doc2.paragraphs)
    for tbl in doc2.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += " ".join("".join(r.text for r in p.runs) for p in cell.paragraphs)

    for old_str in ["000200-50.2019", "RR-000200", "tst_rr_000200"]:
        occ = all_text.count(old_str)
        if occ:
            print(f"[WARN] Post-check: '{old_str}' still present ({occ}×)")
        else:
            print(f"[OK] Post-check: '{old_str}' → 0 occurrences")


if __name__ == "__main__":
    import sys
    dry = "--dry-run" in sys.argv
    main(dry_run=dry)
