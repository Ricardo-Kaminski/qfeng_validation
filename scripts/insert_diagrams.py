"""Insert 5 conceptual diagrams (Diagram 1–5) into the canonical Q-FENG paper DOCX.

For each insertion point:
  1. Find anchor paragraph element
  2. Convert SVG → PNG (cairosvg, DPI=150)
  3. Insert: [anchor] → [prose] → [centered image para] → [Caption para]

Insertions done in REVERSE order to avoid index shifting.
"""
from __future__ import annotations

import datetime
import io
import os
import shutil
import sys
from pathlib import Path

import cairosvg
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

DOCX = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")
WP_DIR = Path("D:/OneDrive/Documentos/- Pós-Doc Q-FENG/Artigos/Onda 1 -  Working Paper arXiv (abr 2026)")
PROSE_MD = Path("artefatos/briefings/DIAGRAM_INSERTION_PROSE.md")

IMG_WIDTH = Inches(5.3)   # ~13.5 cm — fits LNCS text block
SVG_DPI   = 150           # balanced quality / file-size


# ── SVG → PNG conversion ─────────────────────────────────────────────

def svg_to_png(svg_path: Path, dpi: int = SVG_DPI) -> bytes:
    return cairosvg.svg2png(url=str(svg_path), dpi=dpi)


def load_image(svg_path: Path, png_fallback: Path | None = None) -> bytes:
    """Return PNG bytes, using png_fallback if SVG conversion fails."""
    try:
        data = svg_to_png(svg_path)
        print(f"  SVG→PNG: {svg_path.name} ({len(data):,} bytes)")
        return data
    except Exception as exc:
        if png_fallback and png_fallback.exists():
            data = png_fallback.read_bytes()
            print(f"  SVG failed ({exc}); using PNG: {png_fallback.name} ({len(data):,} bytes)")
            return data
        raise


# ── Prose parsing ─────────────────────────────────────────────────────

def parse_prose(md_path: Path) -> dict[str, str]:
    """Extract INS-N prose blocks from DIAGRAM_INSERTION_PROSE.md."""
    text = md_path.read_text(encoding="utf-8")
    sections: dict[str, str] = {}
    current_key = None
    current_lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("## INS-"):
            if current_key and current_lines:
                sections[current_key] = " ".join(
                    l.strip() for l in current_lines if l.strip()
                )
            current_key = line.split()[1]  # "INS-1", "INS-2", etc.
            current_lines = []
        elif line.startswith("---"):
            if current_key and current_lines:
                sections[current_key] = " ".join(
                    l.strip() for l in current_lines if l.strip()
                )
                current_key = None
                current_lines = []
        elif current_key and line.strip() and not line.startswith("#"):
            current_lines.append(line)
    if current_key and current_lines:
        sections[current_key] = " ".join(l.strip() for l in current_lines if l.strip())
    return sections


# ── DOCX XML helpers ──────────────────────────────────────────────────

def find_para_elem(doc: docx.Document, fragment: str):
    """Return the XML element of the first paragraph containing fragment."""
    for p in doc.paragraphs:
        if fragment in p.text:
            return p._element
    raise ValueError(f"Paragraph fragment not found: {fragment!r}")


def make_text_para(doc: docx.Document, text: str, style: str = "Normal",
                   indent: bool = False) -> OxmlElement:
    """Build a paragraph XML element with text and style (not yet inserted)."""
    new_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    style_ids = {
        "Normal": "Normal",
        "Caption": "Legenda",  # styleId in this PT-BR Word document
    }
    pStyle.set(qn("w:val"), style_ids.get(style, style))
    pPr.append(pStyle)
    if indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        pPr.append(ind)
    new_para.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    return new_para


def make_image_para(doc: docx.Document, png_bytes: bytes,
                    width: Inches = IMG_WIDTH) -> OxmlElement:
    """Add centered image at document end (for relationship registration),
    then detach and return the XML element."""
    temp_para = doc.add_paragraph()
    temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = temp_para.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=width)
    # Detach from body
    elem = temp_para._element
    elem.getparent().remove(elem)
    return elem


# ── Main insertion engine ─────────────────────────────────────────────

def insert_diagram(doc: docx.Document,
                   anchor_fragment: str,
                   anchor_is_before: bool,   # True = insert BEFORE anchor para
                   png_bytes: bytes,
                   prose_text: str,
                   caption_text: str,
                   diagram_num: int) -> None:
    """
    Insert [prose] → [image] → [caption] at the correct position.

    anchor_is_before=False: insert AFTER the anchor para
    anchor_is_before=True:  insert BEFORE the anchor para
                            (i.e., after the paragraph before the anchor)
    """
    anchor_elem = find_para_elem(doc, anchor_fragment)

    if anchor_is_before:
        # We need to insert BEFORE anchor_elem.
        # Use anchor_elem's previous sibling as the reference, then addnext.
        parent = anchor_elem.getparent()
        children = list(parent)
        idx = children.index(anchor_elem)
        ref_elem = children[idx - 1] if idx > 0 else None
    else:
        ref_elem = anchor_elem

    # Build elements (prose, image, caption)
    prose_elem   = make_text_para(doc, prose_text, "Normal")
    image_elem   = make_image_para(doc, png_bytes, IMG_WIDTH)
    caption_elem = make_text_para(doc, caption_text, "Caption")

    # Insert in reverse order (each addnext pushes previous down):
    # After ref_elem we want: [prose] [image] [caption]
    # So add caption first, then image, then prose:
    if ref_elem is not None:
        ref_elem.addnext(caption_elem)
        ref_elem.addnext(image_elem)
        ref_elem.addnext(prose_elem)
    else:
        # Edge case: anchor is first element
        parent.insert(0, prose_elem)
        prose_elem.addnext(image_elem)
        image_elem.addnext(caption_elem)

    print(f"  Diagram {diagram_num} inserted (prose + image + caption).")


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")

    # Backup
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = DOCX.with_suffix(f".bak_{ts}.docx")
    shutil.copy2(DOCX, backup)
    print(f"Backup: {backup}")

    # Parse prose
    prose = parse_prose(PROSE_MD)
    print(f"Loaded prose for: {list(prose.keys())}")

    # Pre-convert all SVGs to PNG (do once to catch errors early)
    print("\nConverting SVGs to PNG...")
    png = {
        1: load_image(WP_DIR / "Diagram2_QFENG_Engineering.svg"),
        2: load_image(
            WP_DIR / "Diagram4_Fractal_VSM.svg",
            png_fallback=WP_DIR / "Diagram4_Fractal_VSM.png",
        ),
        3: load_image(WP_DIR / "Diagram1_Interference_Regimes.svg"),
        4: load_image(WP_DIR / "Diagram3_Loss_Landscape.svg"),
        5: load_image(WP_DIR / "Diagram8_Manaus_Timeline.svg"),
    }

    # Load document
    doc = docx.Document(str(DOCX))
    total_paras_before = len(doc.paragraphs)
    print(f"\nDocument loaded: {total_paras_before} paragraphs")

    # ── Define insertion specs ────────────────────────────────────────
    # (anchor_fragment, anchor_is_before, diagram_num, caption)
    insertions = [
        (
            "Three features of this series are theoretically significant",
            True,   # insert BEFORE this paragraph (i.e., right after the Heading)
            5,
            "Diagram 5. Manaus oxygen crisis: actual event sequence (upper track) vs. "
            "contrafactual Q-FENG-mediated response (lower track). The 3-month offset between "
            "the Circuit Breaker activation in October 2020 and the Amazonas state calamity "
            "decree of 23 January 2021 is the governance lead-time quantified by the "
            "Markovian θ_eff formalism.",
        ),
        (
            "when L exceeds the HITL threshold, human review is triggered; "
            "when L exceeds the CB threshold, mandatory intervention is activated",
            False,  # insert AFTER this paragraph
            4,
            "Diagram 4. Conceptual loss landscape of L_Global showing the STAC equilibrium "
            "region. The quantum penalty ridge λ·max(0, −cos θ) excludes non-conforming "
            "configurations from the optimisation trajectory.",
        ),
        (
            # INS-3: anchor = paragraph immediately BEFORE Table 1
            # = Para 103 "where θ ∈ [0°, 180°]... partition this range into three governance regimes"
            "The Q-FENG thresholds partition this range into three governance regimes (Table 1):",
            False,
            3,
            "Diagram 3. Three configurations of preference vectors in the decision Hilbert "
            "space: (a) STAC (constructive interference, cos θ ≈ +1); (b) HITL (intermediate); "
            "(c) CIRCUIT_BREAKER (destructive interference, cos θ ≈ −1).",
        ),
        (
            "Mapping the full Q-FENG C1 architecture to the VSM systems clarifies "
            "the governance role of each pipeline stage",
            False,
            2,
            "Diagram 2. Fractal VSM architecture: scale-invariant S1–S5 structure across "
            "Macro (constitutional/regulatory), Meso (agency-operational), and Micro "
            "(deployment-unit) governance levels.",
        ),
        (
            "This paper presents the Q-FENG (Quantum-Fractal Neurosymbolic Governance) "
            "C1 pipeline",
            False,
            1,
            "Diagram 1. Complete cybernetic inference-audit-feedback cycle of the Q-FENG "
            "architecture. Solid arrows: inference flow. Dashed arrows: logging and data flow. "
            "Long-dashed arrows: feedback channels (Algedonic Signal to S5; Continuous "
            "Training to S1).",
        ),
    ]

    # Execute insertions in order given (they are already in reverse document order,
    # but since we use element references, order doesn't matter for correctness)
    for anchor_frag, before, num, caption in insertions:
        print(f"\n[INS-{num}] Inserting Diagram {num}...")
        insert_diagram(
            doc=doc,
            anchor_fragment=anchor_frag,
            anchor_is_before=before,
            png_bytes=png[num],
            prose_text=prose.get(f"INS-{num}", f"[Diagram {num} connection prose not found]"),
            caption_text=caption,
            diagram_num=num,
        )

    # Save
    doc.save(str(DOCX))
    total_paras_after = len(doc.paragraphs)
    print(f"\nSaved: {DOCX}")
    print(f"Paragraphs: {total_paras_before} → {total_paras_after} (added {total_paras_after - total_paras_before})")

    # File size check
    size_mb = DOCX.stat().st_size / (1024 * 1024)
    print(f"File size: {size_mb:.2f} MB")
    if size_mb > 5.0:
        print("WARNING: File > 5 MB. Consider reducing SVG_DPI.")

    # Post-check
    doc2 = docx.Document(str(DOCX))
    all_text = "\n".join(p.text for p in doc2.paragraphs)

    print("\n=== Post-check ===")

    # Verify 5 Diagrams present
    for n in range(1, 6):
        found = f"Diagram {n}." in all_text
        print(f"  {'OK' if found else 'MISSING'}: Diagram {n} caption")

    # Verify 7 original Figures still present
    for n in range(1, 8):
        found = f"Figure {n}." in all_text
        print(f"  {'OK' if found else 'MISSING'}: Figure {n}")

    # Verify Tables 1-8
    for n in range(1, 9):
        found = f"Table {n}." in all_text
        print(f"  {'OK' if found else 'MISSING'}: Table {n}")

    # Verify Diagrams appear in order in document
    diagram_positions = []
    for i, p in enumerate(doc2.paragraphs):
        for n in range(1, 6):
            if p.text.startswith(f"Diagram {n}.") and p.style.name == "Caption":
                diagram_positions.append((n, i))
    diagram_positions.sort(key=lambda x: x[1])
    nums_in_order = [x[0] for x in diagram_positions]
    print(f"  Diagram order in doc: {nums_in_order}")
    assert nums_in_order == [1, 2, 3, 4, 5], f"Diagrams out of order: {nums_in_order}"

    print("\nAll post-checks passed.")


if __name__ == "__main__":
    main()
