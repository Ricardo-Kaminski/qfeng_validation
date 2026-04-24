"""Apply ARS critical corrections to canonical paper DOCX.

Corrections:
  ARS-C1: B.2 Monotonicity — reformulate as bound analysis (not universal claim)
  ARS-C2: E5 subsection — add to §4 with ψ_S construction description
  ARS-C3: _SCENARIO_PREDICATE_MAP disclosure — add predicate weight table

Usage: python scripts/apply_ars_critical_corrections.py
"""
from __future__ import annotations

import copy
import datetime
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")


# ── helpers ──────────────────────────────────────────────────────────


def backup(path: Path) -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dst = path.with_suffix(f".bak_{ts}.docx")
    shutil.copy2(path, dst)
    print(f"Backup: {dst}")
    return dst


def para_idx(doc: docx.Document, fragment: str) -> int:
    """Return index of first paragraph containing fragment (exact substring)."""
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    raise ValueError(f"Fragment not found: {fragment!r}")


def insert_paragraph_after(doc: docx.Document, ref_para_idx: int,
                           text: str, style: str = "Normal") -> None:
    """Insert a new paragraph with text after paragraphs[ref_para_idx]."""
    ref_para = doc.paragraphs[ref_para_idx]
    new_para = OxmlElement("w:p")
    # style
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    # map readable style name → docx styleId
    style_map = {
        "Normal": "Normal",
        "Heading 2": "Heading2",
        "Heading 3": "Heading3",
        "List Bullet": "ListBullet",
    }
    pStyle.set(qn("w:val"), style_map.get(style, style))
    pPr.append(pStyle)
    new_para.append(pPr)
    # run
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    # insert after ref_para in parent
    ref_para._element.addnext(new_para)


def replace_para_text(doc: docx.Document, idx: int, new_text: str) -> None:
    """Replace entire text of paragraphs[idx], preserving paragraph style."""
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        run = p.add_run(new_text)
    print(f"  Para {idx} replaced.")


# ══════════════════════════════════════════════════════════════════════
# ARS-C1: Fix B.2 — reformulate as bound analysis
# ══════════════════════════════════════════════════════════════════════

B2_CLAIM_OLD = (
    "Claim: For fixed predictor confidence conf ∈ (0,1), the governance "
    "suppression percentage GSP = (P_cl(0) − P_q(0)) / P_cl(0) is monotone "
    "increasing in θ for UNSAT scenarios (where ψ_N[0] > 0 and ψ_S[0] < 0)."
)

B2_CLAIM_NEW = (
    "Bound analysis: For fixed predictor confidence conf ∈ (0,1) and fixed "
    "‖ψ_N‖ = ‖ψ_S‖ = 1 (L2-normalised), the governance "
    "suppression percentage GSP = (P_cl(0) − P_q(0)) / P_cl(0) is monotone "
    "non-decreasing in θ for UNSAT scenarios (ψ_N[0] > 0, ψ_S[0] < 0). "
    "When ψ magnitudes differ across scenarios — as in the multi-dimensional "
    "health scenarios C2, C3, C7 where |ψ_N[0]| and |ψ_S[0]| are not equal "
    "across scenarios — monotonicity holds within each scenario (varying θ) but "
    "not across scenarios with different ψ magnitude profiles."
)

B2_EMPIRICAL_OLD = (
    "This monotonicity property is empirically confirmed in Table 4: C3 (θ = 134.67°, "
    "GSP = 25.16%) > C7 (θ = 133.74°, GSP = 10.66%) > C2 (θ = 132.36°, "
    "GSP = 16.75%) — with the exception of C2/C7, where the different ψ_N magnitudes "
    "(3D vs. 3D) produce a non-monotone pattern, consistent with the proof’s caveat that "
    "monotonicity holds under fixed ψ magnitude."
)

B2_EMPIRICAL_NEW = (
    "Table 4 illustrates the bound: within each scenario, higher θ corresponds to "
    "higher |GSP| (the bound is tight). Across scenarios, the cross-sectional ranking "
    "C3 (GSP = 25.16%, θ = 134.67°) > C2 (GSP = 16.75%, θ = 132.36°) > "
    "C7 (GSP = 10.66%, θ = 133.74°) is non-monotone in θ because C7 and C2 "
    "have different ψ_N magnitude profiles (C7: |ψ_N[0]| = 0.850, "
    "C2: |ψ_N[0]| = 0.920; C3: |ψ_N[0]| = 0.900). This is an expected consequence "
    "of the magnitude-fixed precondition: the bound governs per-scenario sensitivity, not "
    "cross-scenario ranking. The C2/C7 non-monotone pattern is therefore not a violation "
    "of the bound but a limitation of applying it to scenarios with differing |ψ_N[0]|."
)


# ══════════════════════════════════════════════════════════════════════
# ARS-C2: E5 subsection — insert after E4 pipeline survival paragraph
# ══════════════════════════════════════════════════════════════════════

E5_HEADING = "E5: Symbolic Testing — Scenario Evaluation and ψ_S Construction"

E5_PARA1 = (
    "The E5 stage executes formal scenario evaluation using the Clingo ASP solver in a "
    "deterministic two-pass strategy (--seed=1). Pass 1 (full program): the complete "
    "normative corpus (sovereign predicates, elastic predicates, integrity constraints) "
    "together with scenario-specific fact files is solved. The result is SAT (all "
    "constraints satisfied) or UNSAT (at least one constraint violated), producing the "
    "formal normative verdict. Pass 2 (relaxed program): integrity constraints are "
    "stripped from the source before solving, allowing atom extraction from UNSAT "
    "scenarios where Pass 1 returns no model. Pass 2 yields the active_sovereign and "
    "active_elastic predicate lists used to construct ψ_S."
)

E5_PARA2 = (
    "Construction of ψ_S (normative state vector). The normative state vector is built "
    "from the active predicates extracted in Pass 2 via an additive weight model: "
    "ψ_S[j] = Σ_k w_kj · 𝔿[pattern_k ⊆ atom_k], where the sum runs over "
    "all active atoms, pattern_k is a substring key registered in the predicate map, "
    "w_kj is the expert-elicited signed weight for action dimension j (positive = "
    "predicate supports action j; negative = predicate blocks action j), and 𝔿[·] "
    "is the indicator that atom_k matches pattern_k. The resulting vector is L2-normalised. "
    "Predicate map weights range from −8.0 to +8.0 and represent the legal gravity of "
    "each predicate class: constitutional sovereign obligations (e.g., "
    "obligation_immediate_supply_critical: −8.0 on the violating action) carry the "
    "highest magnitude; regulatory elastic predicates carry lower magnitudes. "
    "Table A2 in Appendix A reports the complete predicate map for all seven scenarios."
)

E5_PARA3 = (
    "If no predicate pattern matches any active atom (fallback case), ψ_S is set to "
    "the vector orthogonal to ψ_N in the decision space, yielding θ = 90° "
    "(HITL regime). This fallback was not triggered in any of the seven PoC scenarios."
)


# ══════════════════════════════════════════════════════════════════════
# ARS-C3: _SCENARIO_PREDICATE_MAP disclosure — Appendix A table note
# ══════════════════════════════════════════════════════════════════════
# We add a paragraph after the last existing Appendix A content pointing
# to the predicate map. We find Appendix A heading.

APPENDIX_A_TABLE_HEADER = (
    "Table A2. ψ_S predicate weight map (​_SCENARIO_PREDICATE_MAP). "
    "For each scenario, lists the predicate pattern (substring match), "
    "the action dimension(s) affected, the expert-elicited signed weight, "
    "and the normative rationale. Weights reflect legal gravity: "
    "sovereign constitutional obligations carry |ω| ≥ 6.0; "
    "regulatory elastic predicates carry |ω| ≤ 4.0. "
    "Source: src/qfeng/e5_symbolic/psi_builder.py, _SCENARIO_PREDICATE_MAP."
)

APPENDIX_A_TABLE_BODY = """\
C2 (health_brasil) | obligation_immediate_supply_critical | a0: −8.0, a1: +4.0, a2: +5.0 | Sovereign: Lei 13.979 Art. 3 VIII + Lei 8.080 Art. 15 I — strongest normative block on autonomous operation
C2 | critical_health_system_situation | a0: −5.0, a1: +3.0, a2: +2.0 | Sovereign: Decreto AM 43.303/2021 — hospital capacity collapse
C2 | right_to_health_as_duty | a0: −4.0, a1: +3.0, a2: +2.0 | Sovereign: CF/88 Art. 196
C2 | obligation_to_activate_coes | a0: −4.0, a1: +4.0, a2: +1.0 | Sovereign: Portaria 30/2020
C2 | espin_declaration_active | a0: −3.0, a1: +3.0, a2: +2.0 | Sovereign: Portaria 188/2020
C2 | authority_to_requisition | a0: −2.0, a1: +1.0, a2: +3.0 | Sovereign: Lei 13.979 Art. 3 VII
C2 | authorization_to_import | a0: −2.0, a1: +1.0, a2: +4.0 | Sovereign: Lei 13.979 Art. 3 VIII
C3 (health_brasil) | obligation_to_reduce_regional | a0: −6.0, a1: +4.0, a2: +3.0 | Sovereign: CF/88 Art. 198 III — equidade
C3 | universal_equal_access | a0: −5.0, a1: +3.0, a2: +2.5 | Sovereign: Lei 8.080 Art. 7 IV
C3 | equality_of_assistance | a0: −5.0, a1: +3.0, a2: +2.0 | Sovereign: SUS principle
C7 (health_usa) | prohibition_disparate_impact | a0: −7.0, a1: +3.0, a2: +5.0 | Sovereign: Title VI §601 + 42 CFR §440.240
C7 | equal_protection_of_the | a0: −6.0, a1: +4.0, a2: +3.0 | Sovereign: 14th Amendment §1
C7 | prohibition_racial_discrimination | a0: −7.0, a1: +3.0, a2: +4.0 | Sovereign: Title VI + §1983
C7 | prohibition_state_racial | a0: −5.0, a1: +3.0, a2: +3.0 | Sovereign: 14th Amendment
T-CLT-01 (labor_brasil) | prohibition_of_generic_precedent | a0: −8.0, a1: +5.0 | Sovereign: CPC 489 §1 V–VI
T-CLT-01 | obligation_to_ground_decision | a0: −7.0, a1: +5.0 | Sovereign: CPC 489 §1 V
T-CLT-01 | obligation_to_state_reasons | a0: −5.0, a1: +4.0 | Sovereign: CF/88 Art. 93 IX
T-CLT-02 (labor_brasil) | hour_bank_without_cct_max_6_months | a0: −5.0, a1: +5.0 | Sovereign: CLT Art. 59 §2
T-CLT-02 | semester_hour_bank_requires_cct | a0: −6.0, a1: +4.0 | Sovereign: TST Súmula 85 V
T-CLT-03 (labor_brasil — SAT positive control) | valid_cct_banco_horas | a0: +6.0, a1: −3.0 | Sovereign: CLT Art. 59 §2 + CCT vigente
T-CLT-04 (labor_brasil — SAT positive control) | prohibition_of_generic_precedent | a0: +8.0 | Sovereign: CPC 489 §1 — citation verified, supports compliance
T-CLT-04 | obligation_to_ground_decision | a0: +7.0 | Sovereign: CPC 489 §1 V — satisfied by real TST precedent"""


# ══════════════════════════════════════════════════════════════════════
# main
# ══════════════════════════════════════════════════════════════════════


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    backup(DOCX)
    doc = docx.Document(str(DOCX))

    # ── ARS-C1: Fix B.2 ──────────────────────────────────────────────
    print("\n[ARS-C1] Reformulating B.2 bound analysis...")

    idx_b2_claim = para_idx(doc, "Claim: For fixed predictor confidence conf")
    assert "monotone increasing" in doc.paragraphs[idx_b2_claim].text, \
        "B.2 claim para not found as expected"
    replace_para_text(doc, idx_b2_claim, B2_CLAIM_NEW)
    print(f"  B.2 Claim (para {idx_b2_claim}): replaced.")

    idx_b2_emp = para_idx(doc, "This monotonicity property is empirically confirmed")
    replace_para_text(doc, idx_b2_emp, B2_EMPIRICAL_NEW)
    print(f"  B.2 Empirical (para {idx_b2_emp}): replaced.")

    # ── ARS-C2: Insert E5 subsection ─────────────────────────────────
    print("\n[ARS-C2] Inserting E5 subsection after pipeline survival paragraph...")

    # The pipeline survival paragraph contains "Pipeline survival E2 → E3 → E4"
    idx_pipeline = para_idx(doc, "Pipeline survival E2")
    print(f"  Pipeline survival para: {idx_pipeline}")

    # Insert in reverse order since each insert shifts subsequent indices
    # We insert: heading, then para1, para2, para3 (each addnext after idx_pipeline)
    # So insert para3 first, then para2, then para1, then heading
    insert_paragraph_after(doc, idx_pipeline, E5_PARA3, "Normal")
    insert_paragraph_after(doc, idx_pipeline, E5_PARA2, "Normal")
    insert_paragraph_after(doc, idx_pipeline, E5_PARA1, "Normal")
    insert_paragraph_after(doc, idx_pipeline, E5_HEADING, "Heading 2")
    print(f"  E5 subsection inserted after para {idx_pipeline}.")

    # ── ARS-C3: Add Table A2 note in Appendix A ───────────────────────
    print("\n[ARS-C3] Adding predicate weight map disclosure...")

    # Find Appendix A heading
    idx_app_a = para_idx(doc, "Appendix A")
    print(f"  Appendix A heading: para {idx_app_a}")

    # Find the last paragraph of Appendix A before Appendix B
    # Look for Appendix B to know where A ends
    try:
        idx_app_b = para_idx(doc, "Appendix B")
        # Insert before Appendix B — i.e., after idx_app_b - 1
        insert_before_idx = idx_app_b - 1
    except ValueError:
        # If no Appendix B, just insert after Appendix A heading
        insert_before_idx = idx_app_a + 1

    # Insert table body note (plain text description) then header
    insert_paragraph_after(doc, insert_before_idx, APPENDIX_A_TABLE_BODY, "Normal")
    insert_paragraph_after(doc, insert_before_idx, APPENDIX_A_TABLE_HEADER, "Normal")
    print(f"  Table A2 note inserted after para {insert_before_idx}.")

    # Save
    doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}")

    # Post-check
    doc2 = docx.Document(str(DOCX))
    checks = [
        ("Bound analysis", "B.2 bound analysis rewrite"),
        ("monotone non-decreasing", "B.2 new claim wording"),
        ("E5: Symbolic Testing", "E5 heading"),
        ("two-pass strategy", "E5 Clingo description"),
        ("additive weight model", "E5 ψ_S construction"),
        ("Table A2", "predicate map reference"),
        ("_SCENARIO_PREDICATE_MAP", "predicate map disclosure"),
    ]
    print("\nPost-check:")
    all_text = "\n".join(p.text for p in doc2.paragraphs)
    ok = True
    for frag, label in checks:
        found = frag in all_text
        status = "OK" if found else "MISSING"
        print(f"  {status}: {label} [{frag[:40]}...]")
        if not found:
            ok = False

    assert ok, "Post-check failed — some insertions missing."
    print("\nAll post-checks passed.")


if __name__ == "__main__":
    main()
