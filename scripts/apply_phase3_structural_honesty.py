"""Phase 3 — Structural honesty additions to auditfixed docx.

Inserts after 'Mandatory disclosure 2' (para 309):
  1. Mandatory disclosure 3: EU AI Act / GDPR / Medicaid corpora not evaluated (H-3)
  2. Table: psi_N source classification — 3 columns (H-4)
"""
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from lxml import etree
import copy

DOCX_PATH = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")

DISCLOSURE3_TEXT = (
    "Mandatory disclosure 3 – Regulatory corpus coverage: "
    "The EU AI Act (Regulation (EU) 2024/1689), GDPR (Regulation (EU) 2016/679), "
    "and Medicaid Title XIX (42 USC §1396 et seq.) corpora are fully codified "
    "in the Q-FENG Clingo corpus (eu_ai_act_obligations.lp, gdpr_data_protection.lp, "
    "medicaid_access.lp) but are not evaluated in the current run: no C5, C6, or C8 "
    "scenario is registered in this PoC (C4 LLM predictor integration pending). "
    "The predicates in these files are verified for internal consistency but yield "
    "zero active sovereign atoms in the present results. Paper 2 and future work "
    "will exercise these corpora against planned C5 (EU AI Act high-risk system audit), "
    "C6 (GDPR data minimisation breach), and C8 (Medicaid comparability gap) scenarios."
)

PSI_TABLE_CAPTION = (
    "Table A1. ψ_N Source Classification. "
    "All scenarios in the current PoC use synthetic-calibrated or "
    "pressure-score-interpolated ψ_N vectors. "
    "No predictor-derived ψ_N is evaluated in this paper."
)

PSI_TABLE_HEADERS = ["ψ_N Source Type", "Scenario(s)", "Description"]

PSI_TABLE_ROWS = [
    [
        "pressure-score-interpolated",
        "C2 (12-month series)",
        "Monthly SIH/DATASUS hospital occupancy pressure scores; real institutional "
        "TOH values from FVS-AM (Oct/2020–Mar/2021); epidemiological estimates "
        "for Jul–Sep/2020 and Apr–Jun/2021 (σ=0.10).",
    ],
    [
        "synthetic-calibrated",
        "C2 (single-shot), C3, C7, T-CLT-01, T-CLT-02, T-CLT-03, T-CLT-04",
        "Fixed ψ_N calibrated from literature or normative document counts. "
        "C2 single-shot: FVS-AM boletim 16/jan/2021 (92% UTI occupancy). "
        "C3: regional SUS allocation literature (27 normative docs). "
        "C7: Obermeyer et al. (2019) risk-score disparity (n=48,784). "
        "T-CLT-01–04: normative construction from CLT/CPC corpus.",
    ],
    [
        "predictor-derived",
        "— (none in this PoC)",
        "C4 LLM predictor (Ollama/qwen2.5:14b) integration pending. "
        "Planned for C4a (chain-of-thought grounding), C4b (hallucination), "
        "C4c (adversarial prompt). No predictor-derived ψ_N values are "
        "reported in the current paper.",
    ],
]


def _find_para_index(doc: Document, substring: str) -> int:
    """Return index of first paragraph containing substring, or -1."""
    for i, p in enumerate(doc.paragraphs):
        if substring in p.text:
            return i
    return -1


def _copy_para_style(src_para, dst_para):
    """Copy run formatting (font, size) from first run of src to dst."""
    if src_para.runs and dst_para.runs:
        src_run = src_para.runs[0]
        dst_run = dst_para.runs[0]
        dst_run.font.size = src_run.font.size
        dst_run.font.name = src_run.font.name


def _insert_paragraph_after(doc: Document, ref_para, text: str, bold_prefix: str = "") -> None:
    """Insert a new paragraph after ref_para with given text."""
    new_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    # Copy paragraph style from reference
    if ref_para._p.pPr is not None:
        pPr = copy.deepcopy(ref_para._p.pPr)
    new_para.append(pPr)

    # If there's a bold prefix (like "Mandatory disclosure 3 –"), add it bold
    if bold_prefix and text.startswith(bold_prefix):
        r_bold = OxmlElement("w:r")
        rPr_bold = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr_bold.append(b)
        r_bold.append(rPr_bold)
        t_bold = OxmlElement("w:t")
        t_bold.text = bold_prefix
        t_bold.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_bold.append(t_bold)
        new_para.append(r_bold)

        # Rest of text (non-bold)
        rest = text[len(bold_prefix):]
        r_rest = OxmlElement("w:r")
        t_rest = OxmlElement("w:t")
        t_rest.text = rest
        t_rest.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_rest.append(t_rest)
        new_para.append(r_rest)
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        new_para.append(r)

    ref_para._p.addnext(new_para)
    return new_para


def _insert_table_after_para(doc: Document, ref_para_xml, headers, rows) -> None:
    """Insert a table right after ref_para_xml in the document body."""
    n_cols = len(headers)
    n_rows = 1 + len(rows)  # header + data rows

    tbl = OxmlElement("w:tbl")

    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    tbl.append(tblPr)

    # Table grid
    tblGrid = OxmlElement("w:tblGrid")
    for _ in range(n_cols):
        gridCol = OxmlElement("w:gridCol")
        tbl.append(tblGrid)
    tbl.append(tblGrid)

    def _make_cell(text: str, bold: bool = False) -> OxmlElement:
        tc = OxmlElement("w:tc")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        if bold:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    def _make_row(cells_text, bold=False) -> OxmlElement:
        tr = OxmlElement("w:tr")
        for ct in cells_text:
            tr.append(_make_cell(ct, bold=bold))
        return tr

    # Header row
    tbl.append(_make_row(headers, bold=True))
    # Data rows
    for row in rows:
        tbl.append(_make_row(row))

    # Insert table after ref_para_xml
    ref_para_xml.addnext(tbl)
    return tbl


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"Not found: {DOCX_PATH}")

    bak = DOCX_PATH.with_suffix(f".bak_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    shutil.copy2(DOCX_PATH, bak)
    print(f"Backup: {bak}")

    doc = Document(DOCX_PATH)

    # Find anchor: "Mandatory disclosure 2"
    anchor_idx = _find_para_index(doc, "Mandatory disclosure 2")
    if anchor_idx < 0:
        raise ValueError("Could not find 'Mandatory disclosure 2' paragraph")
    anchor_para = doc.paragraphs[anchor_idx]
    print(f"Found anchor at para {anchor_idx}: '{anchor_para.text[:60]}...'")

    # Step 1: Insert Mandatory disclosure 3 after anchor (para 309)
    bold_prefix = "Mandatory disclosure 3 – Regulatory corpus coverage: "
    _insert_paragraph_after(
        doc,
        anchor_para,
        DISCLOSURE3_TEXT,
        bold_prefix=bold_prefix,
    )
    print("  [OK] Inserted Mandatory disclosure 3")

    # Re-find the new paragraph we just inserted (it's now right after anchor)
    # The disclosure 3 was inserted after anchor_para, so it's at anchor_idx+1
    # But doc.paragraphs might not update in-place — work from XML
    # We need to insert the table caption and table AFTER disclosure 3
    # Disclosure 3 is now the element immediately after anchor_para._p in the body
    disclosure3_xml = anchor_para._p.getnext()

    # Step 2: Insert table caption after disclosure 3
    caption_para = OxmlElement("w:p")
    caption_pPr = OxmlElement("w:pPr")
    caption_para.append(caption_pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    i_elem = OxmlElement("w:i")
    rPr.append(i_elem)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = PSI_TABLE_CAPTION
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    caption_para.append(r)
    disclosure3_xml.addnext(caption_para)
    print("  [OK] Inserted table caption")

    # Step 3: Insert table after caption
    _insert_table_after_para(doc, caption_para, PSI_TABLE_HEADERS, PSI_TABLE_ROWS)
    print("  [OK] Inserted psi_N source classification table (3 cols x 4 rows including header)")

    doc.save(DOCX_PATH)
    print(f"\nSaved: {DOCX_PATH}")

    # Post-check
    doc2 = Document(DOCX_PATH)
    checks = [
        "Mandatory disclosure 3",
        "EU AI Act",
        "predictor-derived",
        "pressure-score-interpolated",
        "synthetic-calibrated",
        "Table A1",
    ]
    for check in checks:
        found = any(check in p.text for p in doc2.paragraphs)
        # Also check tables
        if not found:
            for tbl in doc2.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if check in cell.text:
                            found = True
        print(f"  {'[OK]' if found else '[WARN]'} '{check}': {'present' if found else 'MISSING'}")


if __name__ == "__main__":
    main()
