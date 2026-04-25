"""
Generates Table 1 (S2.8 jurisdictional fractal mapping) as native Word table.
Output: TABLE1_S2-8_Fractal_Mapping.docx — open, copy table, paste into paper.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"C:\Workspace\academico\qfeng_validacao\artefatos\briefings\TABLE1_S2-8_Fractal_Mapping.docx")

doc = Document()

# Page setup — landscape would help but keep portrait for paper compatibility
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Title
p = doc.add_paragraph()
run = p.add_run("Table 1 — Q-FENG ↔ VSM mapping across the jurisdictional stack")
run.bold = True
run.font.size = Pt(11)

# Caption
p = doc.add_paragraph()
run = p.add_run(
    "Institutional mapping of Viable System Model functions across the three fractal levels of "
    "Q-FENG governance. Brazilian institutional examples are given as primary references; "
    "analogous bodies operate in the EU and US regimes covered by the Q-FENG validation. "
    "The Q-FENG framework provides computational instantiation of the S2+S3+S3* triad at the "
    "Micro level (highlighted in bold), with semantic propagation to higher levels through "
    "formal derivability of the jurisdictional stack."
)
run.italic = True
run.font.size = Pt(9)

doc.add_paragraph()  # spacing

# Build table: 7 rows (1 header + 6 systems) × 4 cols (System / Macro / Meso / Micro)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths (total ~16cm portrait)
widths = [Cm(2.5), Cm(4.5), Cm(4.5), Cm(4.5)]
for i, w in enumerate(widths):
    for row in table.rows:
        row.cells[i].width = w

def set_cell(cell, text, bold=False, italic=False, size=9, fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

def set_micro_cell(cell, prefix_text, qfeng_text, suffix_text='', size=9):
    """Micro cell with Q-FENG component in bold."""
    cell.text = ''
    p = cell.paragraphs[0]
    if prefix_text:
        r0 = p.add_run(prefix_text)
        r0.font.size = Pt(size)
    r1 = p.add_run(qfeng_text)
    r1.font.size = Pt(size)
    r1.bold = True
    if suffix_text:
        r2 = p.add_run(suffix_text)
        r2.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# Header row
hdr = table.rows[0].cells
set_cell(hdr[0], "VSM System", bold=True, size=10, fill='2E5C8A')
set_cell(hdr[1], "Macro (Constitutional)", bold=True, size=10, fill='2E5C8A')
set_cell(hdr[2], "Meso (Sectoral / Infralegal)", bold=True, size=10, fill='2E5C8A')
set_cell(hdr[3], "Micro (Algorithmic / Operational)", bold=True, size=10, fill='2E5C8A')

# Set header text color to white
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Row 6 — S1 Operations (bottom row: S5→S1 top-to-bottom order)
r = table.rows[6].cells
set_cell(r[0], "S1\nOperations", bold=True)
set_cell(r[1], "National-scale State operations: federal programmes, executive administrative acts (federal ministries, autarchies; in EU: Commission directorates; in US: federal agencies)")
set_cell(r[2], "Sectoral execution: state health secretariats, regional INSS units, regional labour tribunals (in EU: national health authorities; in US: state Medicaid programmes)")
set_micro_cell(r[3], "", "Algorithmic predictor producing ψ_N", " (LightGBM, time-series, ASP rule engine, LLM); or institutional deployment unit")

# Row 5 — S2 Coordination
r = table.rows[5].cells
set_cell(r[0], "S2\nCoordination", bold=True)
set_cell(r[1], "Federative pact, intergovernmental cooperation laws, supreme court binding precedent in uniformising function (CNJ, intergovernmental commissions; in EU: Council coordination; in US: federal-state agreements)")
set_cell(r[2], "Sectoral coordination: bipartite intergestor commissions (CIB), technical chambers, sectoral clinical protocols")
set_micro_cell(r[3], "", "Cross-corpus consistency layer", ": concurrency-pair detection at E1 (Jaccard ≥ 0.55), SHA-256 caching; deployment-unit operational protocols")

# Row 4 — S3 Operational Control
r = table.rows[4].cells
set_cell(r[0], "S3\nOperational Control", bold=True)
set_cell(r[1], "Direct command authority: Presidency, Ministers, top federal executive leadership (in EU: Commission College; in US: Cabinet-level command)")
set_cell(r[2], "Agency direction: state Secretaries, autarchy presidents, sectoral managers with direct authority")
set_micro_cell(r[3], "", "Circuit Breaker logic", " (θ ≥ 120° threshold suspending autonomous operation); deployment-unit management")

# Row 3 — S3* Audit Channel
r = table.rows[3].cells
set_cell(r[0], "S3*\nAudit Channel", bold=True, fill='FFF8E7')
set_cell(r[1], "Constitutional-level audit: TCU, MPF, STF in concentrated control, organised social control (in EU: Court of Auditors, Ombudsman; in US: GAO, Inspectors General)", fill='FFF8E7')
set_cell(r[2], "Sectoral audit: DENASUS in SUS, internal controls, sectoral ombudsmen, state audit courts (TCEs); regulatory agency oversight", fill='FFF8E7')
# Custom Micro cell with bold + fill
r[4-1].text = ''  # already at index 3; this line is just defensive
cell = r[3]
cell.text = ''
p = cell.paragraphs[0]
r1 = p.add_run("Clingo SAT/UNSAT evaluation + Pass 2 active sovereign predicate analysis")
r1.font.size = Pt(9)
r1.bold = True
r2 = p.add_run("; deployment-unit ombudsman and local social control")
r2.font.size = Pt(9)
cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
tcPr = cell._tc.get_or_add_tcPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'FFF8E7')
tcPr.append(shd)

# Row 2 — S4 Intelligence
r = table.rows[2].cells
set_cell(r[0], "S4\nIntelligence & Adaptation", bold=True)
set_cell(r[1], "National prospective intelligence: IPEA, IBGE, federal observatories, governmental BI, Casa Civil strategic advisory, theory of change applied to national policy (in EU: Joint Research Centre; in US: GAO and CBO analytical units)")
set_cell(r[2], "Sectoral BI: agency technical areas, sectoral observatories (e.g. Observatório SUS), impact evaluation, technical strategic advisory")
set_micro_cell(r[3], "", "Markovian θ_eff with adaptive memory α(t)", "; deployment-unit monitoring dashboards")

# Row 1 — S5 Policy & Identity (top row)
r = table.rows[1].cells
set_cell(r[0], "S5\nPolicy & Identity", bold=True)
set_cell(r[1], "Federal Constitution of 1988, petreous clauses (Art. 60 §4º), constitutional treaties, STF binding jurisprudence (in EU: TFEU + Charter of Fundamental Rights; in US: Constitution + Reconstruction Amendments)")
set_cell(r[2], "Sectoral organic laws (Lei 8.080/SUS, CLT, Lei 13.467/2017), structural ministerial portarias, regulatory agency rules, sectoral decrees")
set_micro_cell(r[3], "", "Clingo predicates derived through E0–E5 pipeline", ": sovereign + elastic predicates with HITL-validated sovereignty classification")

# Set borders for clean look
def set_table_borders(tbl):
    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._element.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tblBorders.append(b)
    tblPr.append(tblBorders)

set_table_borders(table)

# Save
doc.save(str(OUT))
print(f"Generated: {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
