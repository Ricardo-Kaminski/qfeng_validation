"""Apply literato rewrites and NIST AI RMF paragraph to canonical Q-FENG paper.

Changes:
  LIT-1: §1 opening hook — leads with concrete Manaus + Obermeyer crises (was generic framing)
  LIT-2: §6.4 ablation intro — honest framing of parity result, not defensive
  LIT-3: §6.4 ablation conclusion — sharp two-question structure
  LIT-4: §8 closing paragraph — memorable final sentence
  NIST-1: §7.1 — NIST AI RMF MEASURE function positioning paragraph
"""
from __future__ import annotations

import datetime
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")


def backup(path: Path) -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dst = path.with_suffix(f".bak_{ts}.docx")
    shutil.copy2(path, dst)
    print(f"Backup: {dst}")
    return dst


def para_idx(doc: docx.Document, fragment: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    raise ValueError(f"Fragment not found: {fragment!r}")


def replace_para_text(doc: docx.Document, idx: int, new_text: str) -> None:
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"  Para {idx} replaced.")


def insert_paragraph_after(doc: docx.Document, ref_idx: int,
                            text: str, style_id: str = "Normal") -> None:
    ref_para = doc.paragraphs[ref_idx]
    new_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    new_para.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    ref_para._element.addnext(new_para)
    print(f"  Paragraph inserted after {ref_idx}.")


# ══════════════════════════════════════════════════════════════════════
# LIT-1: §1 opening hook
# ══════════════════════════════════════════════════════════════════════

S1_HOOK_OLD = (
    "The governance of algorithmic decision systems in public administration poses "
    "a challenge that transcends the standard paradigm of statistical performance "
    "evaluation. A system may exhibit high predictive accuracy by conventional metrics "
    "while simultaneously violating constitutional mandates, circumventing legally "
    "protected rights, or allocating resources through mechanisms structurally "
    "prohibited by statute. The canonical example is the healthcare algorithm analysed "
    "by Obermeyer et al. (2019): deployed across hundreds of US hospitals, the system "
    "exhibited a 34-percentage-point racial gap in enrolment recommendations for the "
    "same health need — a failure that standard performance dashboards did not detect, "
    "and that only became visible through normative external audit comparing actual "
    "outputs against the equal-protection principles encoded in the 14th Amendment §1 "
    "Equal Protection Clause and Title VI of the Civil Rights Act 1964 (42 U.S.C. §2000d)."
)

S1_HOOK_NEW = (
    "On 23 January 2021, the state of Amazonas declared public calamity as Manaus "
    "hospitals ran out of oxygen and patients died unventilated. The normative "
    "architecture for responding had been in force for decades — CF/88 Art. 196 "
    "declaring health a fundamental right, Lei 8.080/1990 Art. 7 mandating universal "
    "equitable access, Lei 13.979/2020 Art. 3 VIII authorising emergency requisition — "
    "but no formal mechanism existed to measure the widening gap between statutory "
    "obligation and institutional execution. Three months earlier, ICU occupancy had "
    "already crossed into CIRCUIT_BREAKER territory by the metric developed in this "
    "paper. The healthcare algorithm studied by Obermeyer et al. (2019) represents the "
    "complementary failure: deployed across hundreds of US hospitals, it produced a "
    "34-percentage-point racial gap in enrolment recommendations for the same health "
    "need — violating the equal-protection mandate of the 14th Amendment §1 Equal "
    "Protection Clause and Title VI (42 U.S.C. §2000d) in a way that standard "
    "performance dashboards did not detect because they measured statistical accuracy, "
    "not normative alignment. Two failure modes, two jurisdictions, one diagnostic gap. "
    "This paper addresses that gap."
)

# ══════════════════════════════════════════════════════════════════════
# LIT-2: §6.4 ablation intro — honest framing
# ══════════════════════════════════════════════════════════════════════

ABLATION_INTRO_OLD = (
    "The rule-based baseline classifies all 7 scenarios correctly (7/7, 100%), "
    "matching Q-FENG's classification on every case. This confirms the reviewer "
    "concern that for the classification task alone, the quantum formalism is not "
    "strictly necessary: a predicate counter achieves identical results."
)

ABLATION_INTRO_NEW = (
    "The result is unambiguous: the rule-based baseline classifies all 7 scenarios "
    "correctly (7/7, 100%), matching Q-FENG's classification on every case. For binary "
    "regime classification alone, the quantum formalism is not strictly necessary — "
    "a predicate counter achieves identical results. We report this finding directly "
    "rather than minimise it: classification parity with a simpler baseline is a "
    "genuine result, and suppressing it would misrepresent the contribution."
)

# ══════════════════════════════════════════════════════════════════════
# LIT-3: §6.4 ablation conclusion — two-question structure
# ══════════════════════════════════════════════════════════════════════

ABLATION_CONCL_OLD = (
    "The ablation thus supports an important qualification: Q-FENG is not proposed "
    "as a more accurate classifier than rule-based approaches (classification accuracy "
    "is already at ceiling for both). It is proposed as a richer governance measurement "
    "instrument that quantifies degree, direction, and temporal evolution of normative "
    "misalignment — capabilities that rule-based predicate counting cannot provide."
)

ABLATION_CONCL_NEW = (
    "The ablation draws a sharp line between two questions that must not be conflated. "
    "First: can governance regime classification be performed without quantum "
    "mathematics? Yes — a predicate counter achieves it equally well. Second: does "
    "Q-FENG add governance measurement capabilities beyond regime classification? "
    "Yes — three of them, all demonstrated in this PoC and none available from any "
    "predicate-counting approach. The Q-FENG proposition concerns the second question. "
    "Classifying CIRCUIT_BREAKER is a starting point; quantifying by how much the "
    "interference geometry suppresses violation probability, and tracking how that "
    "quantity evolves through a crisis trajectory, are the capabilities that the "
    "governance use case requires and that the quantum formalism uniquely provides."
)

# ══════════════════════════════════════════════════════════════════════
# NIST-1: NIST AI RMF MEASURE positioning (insert after cosine-similarity para)
# ══════════════════════════════════════════════════════════════════════

NIST_PARA = (
    "Q-FENG maps directly to the NIST Artificial Intelligence Risk Management "
    "Framework (AI RMF 1.0; NIST, 2023) MEASURE function. The AI RMF specifies that "
    "AI risks should be quantified, tracked, and reported against established metrics "
    "(MEASURE 2.5, MEASURE 2.6, MEASURE 4.2), but provides no computational "
    "specification for normative alignment measurement in particular. Q-FENG supplies "
    "that specification: the interference angle θ implements MEASURE 2.5's call for "
    "vulnerability metrics (governance misalignment as angular distance in Hilbert "
    "space); the governance suppression percentage GSP implements MEASURE 2.6's call "
    "for quantitative risk-impact assessment (the probability mass by which the quantum "
    "model suppresses norm-violating actions relative to a classical baseline); and the "
    "Markovian θ_eff recurrence implements MEASURE 4.2's directive that measurement "
    "results inform evolving risk management decisions. The CIRCUIT_BREAKER activation "
    "at θ ≥ 120° maps to the MANAGE function's response planning tier (MS-2.2), "
    "providing the computational circuit that closes the loop between MEASURE and "
    "MANAGE. Q-FENG is therefore not a supplement to the NIST AI RMF — it is a formal "
    "instantiation of the normative alignment measurement apparatus that the MEASURE "
    "function presupposes but does not specify."
)

# ══════════════════════════════════════════════════════════════════════
# LIT-4: §8 closing paragraph
# ══════════════════════════════════════════════════════════════════════

S8_CLOSE_OLD = (
    "The Q-FENG framework is not a substitute for regulatory text; it is the formal "
    "infrastructure that the EU AI Act presupposes without specifying — the "
    "computational layer at which its Article 14 human oversight obligations and "
    "Article 9 risk management requirements can be given executable content. The "
    "combination of rigorous ASP-based normative evaluation, quantum interference "
    "geometry, and Markovian temporal tracking provides a transition from post-hoc "
    "audit to continuous prospective governance — the standard that high-stakes AI "
    "systems in health, welfare, and legal adjudication require."
)

S8_CLOSE_NEW = (
    "The Q-FENG framework is not a substitute for regulatory text; it is the missing "
    "computational layer between the EU AI Act's requirements and their operational "
    "enforcement. Article 9 mandates continuous risk management; Article 14 mandates "
    "human oversight at the circuit-breaker threshold; Article 15 mandates accuracy "
    "and robustness. Q-FENG gives each of those mandates an executable form: the "
    "interference angle θ that fires the circuit breaker, the governance suppression "
    "percentage that quantifies suppressed violation risk, and the Markovian θ_eff "
    "recurrence that tracks governance trajectories across time. Between the formal "
    "obligation and the algorithmic act, there is a calculable distance. "
    "Measuring it is what Q-FENG does."
)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    backup(DOCX)
    doc = docx.Document(str(DOCX))

    # ── LIT-1: §1 hook ───────────────────────────────────────────────
    print("\n[LIT-1] Rewriting §1 opening hook...")
    idx = para_idx(doc, "The governance of algorithmic decision systems in public administration")
    print(f"  §1 opening at para {idx}")
    replace_para_text(doc, idx, S1_HOOK_NEW)

    # ── LIT-2: §6.4 ablation intro ───────────────────────────────────
    print("\n[LIT-2] Rewriting §6.4 ablation intro...")
    idx = para_idx(doc, "The rule-based baseline classifies all 7 scenarios correctly")
    print(f"  Ablation intro at para {idx}")
    replace_para_text(doc, idx, ABLATION_INTRO_NEW)

    # ── LIT-3: §6.4 ablation conclusion ─────────────────────────────
    print("\n[LIT-3] Rewriting §6.4 ablation conclusion...")
    idx = para_idx(doc, "The ablation thus supports an important qualification")
    print(f"  Ablation conclusion at para {idx}")
    replace_para_text(doc, idx, ABLATION_CONCL_NEW)

    # ── NIST-1: NIST AI RMF paragraph ────────────────────────────────
    print("\n[NIST-1] Inserting NIST AI RMF MEASURE paragraph in §7.1...")
    # Insert after the Fairlearn/IBM comparison sentence (last para before §7.2)
    idx = para_idx(doc, "a broader comparison against Fairlearn, IBM AI Fairness 360")
    print(f"  Fairlearn para at {idx}")
    insert_paragraph_after(doc, idx, NIST_PARA, "Normal")

    # ── LIT-4: §8 closing paragraph ──────────────────────────────────
    print("\n[LIT-4] Rewriting §8 closing paragraph...")
    idx = para_idx(doc, "The Q-FENG framework is not a substitute for regulatory text")
    print(f"  §8 closing at para {idx}")
    replace_para_text(doc, idx, S8_CLOSE_NEW)

    doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}")

    # Post-check
    doc2 = docx.Document(str(DOCX))
    all_text = "\n".join(p.text for p in doc2.paragraphs)
    checks = [
        ("On 23 January 2021, the state of Amazonas", "LIT-1: Manaus hook"),
        ("Two failure modes, two jurisdictions, one diagnostic gap", "LIT-1: hook conclusion"),
        ("We report this finding directly rather than minimise it", "LIT-2: honest framing"),
        ("The ablation draws a sharp line between two questions", "LIT-3: two-question structure"),
        ("Classifying CIRCUIT_BREAKER is a starting point", "LIT-3: conclusion"),
        ("NIST Artificial Intelligence Risk Management", "NIST-1: NIST paragraph"),
        ("MEASURE 2.5", "NIST-1: MEASURE 2.5"),
        ("MEASURE 4.2", "NIST-1: MEASURE 4.2"),
        ("Measuring it is what Q-FENG does", "LIT-4: closing line"),
        ("Between the formal obligation and the algorithmic act", "LIT-4: penultimate line"),
    ]
    print(f"\nTotal paragraphs: {len(doc2.paragraphs)}")
    print("\nPost-check:")
    ok = True
    for frag, label in checks:
        found = frag in all_text
        status = "OK" if found else "MISSING"
        print(f"  {status}: {label}")
        if not found:
            ok = False
    assert ok, "Post-check failed."
    print("\nAll checks passed.")


if __name__ == "__main__":
    main()
