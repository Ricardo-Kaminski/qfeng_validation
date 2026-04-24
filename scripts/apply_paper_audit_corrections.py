"""Phase 2 paper prose corrections — pre-submission audit (UGR review).

Applies to: docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo.docx
Outputs:    docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx

Audit items applied (F0-independent):
  C-4 : Portaria GM/MS 69/2021 → correct normative anchors
  H-1 : SSA §1902(a)(19) → 14th Amendment + Title VI (Obermeyer anchor)
  H-2 : CLT Art. 59-B §1 → CLT Art. 59 §§2 e 5 + Art. 611-A I
  H-4 : LightGBM trained → synthetic calibrated (C3 psi_N)
  H-8 : Table 7 data + narrative → current pipeline values (Sep/2020 peak)

NOT applied (require Phase 0 results):
  C-6 : TST-RR-000200-50.2019.5.02.0020 file (pending F0-1)
  H-5 : Portaria 268/2021 anchor (pending F0-2)

Usage:
    python scripts/apply_paper_audit_corrections.py [--dry-run]
"""

from __future__ import annotations

import argparse
import csv
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# ── Paths ──────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
PAPERS = ROOT / "docs" / "papers"
INPUT_DOCX = PAPERS / "PAPER1_QFENG_FINAL_prob_dados_clingo.docx"
OUTPUT_DOCX = PAPERS / "PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx"
TABLE7_CSV = ROOT / "outputs" / "table7_new_values.csv"

# ── Month label helpers ─────────────────────────────────────────────────────
_PT_MONTHS = ["jan", "fev", "mar", "abr", "mai", "jun",
               "jul", "ago", "set", "out", "nov", "dez"]

def _competencia_label(row: dict) -> str:
    m = int(row["mes_cmpt"])
    y = int(row["ano_cmpt"])
    return f"{_PT_MONTHS[m - 1]}/{y}"

# ── Text replacement in docx ────────────────────────────────────────────────

def _replace_in_run(run, old: str, new: str) -> int:
    """Replace text within a single run. Returns 1 if replaced, else 0."""
    if old in run.text:
        run.text = run.text.replace(old, new)
        return 1
    return 0


def _replace_in_paragraph(para, old: str, new: str) -> int:
    """Replace across all runs of a paragraph. Returns count of replacements."""
    # Check if the full string spans multiple runs (combine and replace)
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return 0
    count = 0
    # Try simple per-run replacement first
    for run in para.runs:
        count += _replace_in_run(run, old, new)
    if count > 0:
        return count
    # Fallback: rebuild paragraph text in first run, clear others
    # (loses per-run formatting but preserves paragraph style)
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]:
            run.text = ""
        return 1
    return 0


def apply_replacements(doc: Document, replacements: list[tuple[str, str]],
                       dry_run: bool = False) -> dict[str, int]:
    """Apply a list of (old, new) replacements across all paragraphs and tables.

    Returns dict mapping old_text → count of replacements made.
    """
    counts: dict[str, int] = {old: 0 for old, _ in replacements}

    def _process_para(para):
        for old, new in replacements:
            n = _replace_in_paragraph(para, old, new)
            counts[old] += n

    for para in doc.paragraphs:
        _process_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)

    return counts

# ── Table 7 update ──────────────────────────────────────────────────────────

def _load_table7_csv() -> list[dict]:
    with open(TABLE7_CSV, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def _find_table7(doc: Document) -> int | None:
    """Return index of the Manaus time-series table, or None."""
    for i, table in enumerate(doc.tables):
        for row in table.rows:
            texts = [c.text.strip() for c in row.cells]
            # Look for rows with "jan/2021" or "jul/2020" in first cell
            if any(t in ("jan/2021", "jul/2020", "set/2020") for t in texts):
                return i
    return None


def update_table7(doc: Document, rows_csv: list[dict],
                  dry_run: bool = False) -> list[str]:
    """Update Table 7 rows with current pipeline values.

    Returns list of log messages.
    """
    logs = []
    idx = _find_table7(doc)
    if idx is None:
        logs.append("WARNING: Table 7 (Manaus series) not found in document")
        return logs

    table = doc.tables[idx]
    logs.append(f"Found Table 7 at table index {idx}")

    # Build a mapping from month label → CSV row
    csv_map = {_competencia_label(r): r for r in rows_csv}

    # Header row detection: skip rows where first cell looks like a header
    for table_row in table.rows:
        cells = table_row.cells
        if not cells:
            continue
        label = cells[0].text.strip()
        if label not in csv_map:
            continue
        r = csv_map[label]

        # Expected columns (0-indexed): label|theta_t|theta_eff|alpha|regime|occ|source|ci_lo|ci_hi
        updates = [
            (1, f"{float(r['theta_t']):.2f}"),
            (2, f"{float(r['theta_efetivo']):.2f}"),
            (3, f"{float(r['alpha_t']):.3f}"),
            (4, r["interference_regime"]),
            (5, f"{int(r['hospital_occupancy_pct'])}%"),
            (6, "SIH/DATASUS"),
            (7, f"{float(r['theta_ci_lower_95']):.2f}"),
            (8, f"{float(r['theta_ci_upper_95']):.2f}"),
        ]

        old_vals = [c.text.strip() for c in cells]
        for col_idx, new_val in updates:
            if col_idx < len(cells):
                cell = cells[col_idx]
                old = cell.text.strip()
                if old != new_val:
                    logs.append(f"  [{label}] col{col_idx}: '{old}' → '{new_val}'")
                    if not dry_run:
                        # Clear cell and set new text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = ""
                        if cell.paragraphs:
                            cell.paragraphs[0].runs[0].text = new_val if cell.paragraphs[0].runs else None
                            if not cell.paragraphs[0].runs:
                                cell.paragraphs[0].add_run(new_val)
                        else:
                            cell.add_paragraph(new_val)

    return logs

# ── Replacement tables ──────────────────────────────────────────────────────

# C-4: Portaria 69/2021 replacements
# Each tuple is (old_text, new_text). Order matters — more specific first.
C4_REPLACEMENTS = [
    # Predicate table: specific anchor reference
    (
        "emergency_obligation_coes/2 | Portaria 69/2021 Art. 1 | SOVEREIGN",
        "emergency_obligation_coes/2 | Decreto AM 43.303/2021 + Lei 13.979/2020 Art.3 VIII | SOVEREIGN",
    ),
    # Table 2 corpus row
    (
        "Portaria 69/2021 + Lei 13.979/2020",
        "Decreto AM 43.303/2021 + Lei 13.979/2020 Art.3 VIII",
    ),
    # Narrative — certified by Portaria
    (
        "certified by Portaria MS 69/2021, Art. 1, §1",
        "documented by FVS-AM Boletim Epidemiológico 16/jan/2021 (103.7% UTI occupancy)",
    ),
    # Time series narrative — declared by
    (
        "declared by Portaria 69/2021",
        "documented by FVS-AM Boletim Epidemiológico 16/jan/2021",
    ),
    (
        "declared by Portaria MS 69/2021",
        "documented by FVS-AM Boletim Epidemiológico 16/jan/2021",
    ),
    # Introduction / normative architecture mention
    (
        "CF/88 Art. 196; Lei 8.080/1990 Art. 7; Portaria 69/2021",
        "CF/88 Art. 196; Lei 8.080/1990 Art. 7; Lei 13.979/2020 Art.3 VIII",
    ),
    # Generic fallback — any remaining mention
    ("Portaria MS 69/2021", "Decreto AM 43.303/2021"),
    ("Portaria GM/MS 69/2021", "Decreto AM 43.303/2021"),
    ("Portaria 69/2021", "Decreto AM 43.303/2021"),
]

# H-1: §1902(a)(19) corrections
H1_REPLACEMENTS = [
    # Main empirical anchor sentence (introduction)
    (
        "equal-protection principles encoded in §1902(a)(19) of the US Social Security Act",
        "equal-protection principles encoded in the 14th Amendment §1 Equal Protection Clause and Title VI of the Civil Rights Act 1964 (42 U.S.C. §2000d)",
    ),
    # C7 scenario anchor sentence
    (
        "encodes the equal-protection obligation of §1902(a)(19) of the US Social Security Act as a sovereign predicate",
        "encodes the equal-protection obligation of the 14th Amendment §1 EPC and Title VI §601 (42 U.S.C. §2000d) as sovereign predicates",
    ),
    # Closing paragraph — CF/88 + SSA closing sentence
    (
        "CF/88 Art. 196 (universal right to health) and SSA §1902(a)(19) (best-interest standard)",
        "CF/88 Art. 196 (universal right to health) and the 14th Amendment §1 EPC (equal protection)",
    ),
    # Listing 1 label
    (
        "USA — SSA §1902(a)(19):",
        "USA — 14th Amendment §1 EPC + Title VI (42 U.S.C. §2000d):",
    ),
    # Listing 1 caption
    (
        "(c) USA — SSA §1902(a)(19).",
        "(c) USA — 14th Amendment §1 EPC + Title VI §601.",
    ),
    # Clingo derivation chain (C7 code block)
    (
        "eligible_individuals) (SSA §1902(a)(19));",
        "eligible_individuals) (14th Amend. §1 + 42 U.S.C. §2000d);",
    ),
    # Predicate table
    (
        "best_interest_standard/2 | SSA §1902(a)(19) | SOVEREIGN",
        "equal_protection_14th/2 | 14th Amendment §1 EPC + 42 U.S.C. §2000d | SOVEREIGN",
    ),
    # Fallback — catches predicate table cell and any remaining standalone reference
    ("SSA §1902(a)(19)", "14th Amendment §1 EPC + 42 U.S.C. §2000d"),
]

# H-2: CLT Art. 59-B corrections
H2_REPLACEMENTS = [
    # Predicate table
    (
        "collective_bargaining_required/2 | CLT Art. 59-B §1 | SOVEREIGN",
        "collective_bargaining_required/2 | CLT Art. 59 §§2 e 5 + Art. 611-A I | SOVEREIGN",
    ),
    # Labour Reform quote (wrong attribution)
    (
        'which introduced CLT Art. 59-B §1: "The establishment of the time account for exceeding forty hours of weekly work may only occur by collective labour convention or collective bargaining agreement."',
        "which introduced CLT Art. 59 §§2 and 5 and Art. 611-A I: the annual bank-of-hours may only be established by collective bargaining (CCT), while the semester bank requires at minimum an agreement (ACE).",
    ),
    # Corpus description
    (
        "CLT Art. 59 and Art. 59-B (hour bank, added by Lei 13.467/2017 Labour Reform)",
        "CLT Art. 59 §§2 and 5 and Art. 611-A I (hour bank collective bargaining requirements, Lei 13.467/2017)",
    ),
    # Fallback — catches predicate table cell and any remaining standalone reference
    ("CLT Art. 59-B §1", "CLT Art. 59 §§2 e 5 + Art. 611-A I"),
]

# H-4: LightGBM training claim (C3)
H4_REPLACEMENTS = [
    (
        "The LightGBM predictor — trained on the normative document count per regional administrative unit across all 27 corpus documents — assigns high weight to the concentrated metropolitan pattern",
        "The ψ_N vector for C3 — calibrated from literature-documented regional SUS allocation patterns (synthetic calibration from 27 corpus normative documents; see Table 3) — assigns high weight to the concentrated metropolitan pattern",
    ),
]

# H-8: Abstract and narrative — peak month correction
H8_NARRATIVE_REPLACEMENTS = [
    # Abstract
    (
        "The Manaus theta-efetivo series reached its peak at θ_eff = 130.85° in February 2021, with Circuit Breaker first activating in October 2020.",
        "The Manaus theta-efetivo series reached its peak at θ_eff = 130.91° in September 2020, with Circuit Breaker first activating in July 2020 and persisting through the 12-month series.",
    ),
    # Series analysis paragraph — "peak in February 2021"
    (
        "Memory-dampened recovery: After the peak in February 2021 (θ_eff = 130.85°)",
        "Memory-dampened recovery: After the peak in September 2020 (θ_eff = 130.91°)",
    ),
    # Bootstrap CI paragraph — Jan/2021 as "peak month"
    (
        "The January 2021 peak month has CI [126.72°, 131.22°] — entirely within the CIRCUIT_BREAKER regime regardless of bootstrap variation.",
        "The September 2020 peak month has CI [130.63°, 132.84°] — entirely within the CIRCUIT_BREAKER regime regardless of bootstrap variation. January 2021 (θ_eff = 118.73°, HITL) reflects the Markovian memory dampening from the preceding December 2020 dip.",
    ),
    # VSM paragraph — "before ICU system exceeded capacity in January 2021"
    (
        "before the ICU system exceeded capacity in January 2021.",
        "before the ICU system's maximum occupancy in late 2020.",
    ),
    # CB activation in Fig 3 caption
    (
        "Circuit Breaker first activates in October 2020 (θ_eff = 125.34°)",
        "Circuit Breaker first activates in July 2020 (θ_eff = 124.88°)",
    ),
    (
        "Peak February 2021 at θ_eff = 130.9°",
        "Peak September 2020 at θ_eff = 130.91°",
    ),
    # Conclusions bullet
    (
        "CIRCUIT_BREAKER activation in October 2020 — three months before the Portaria 69/2021 ICU collapse declaration",
        "CIRCUIT_BREAKER activation from July 2020 — with peak at September 2020 (θ_eff = 130.91°)",
    ),
    # Series CB onset value
    (
        "Circuit Breaker first activates in October 2020 (θ_eff = 125.34°), when occupancy reaches 72%",
        "Circuit Breaker first activates in July 2020 (θ_eff = 124.88°), with the peak in September 2020 (θ_eff = 130.91°)",
    ),
]

ALL_REPLACEMENTS = (
    C4_REPLACEMENTS
    + H1_REPLACEMENTS
    + H2_REPLACEMENTS
    + H4_REPLACEMENTS
    + H8_NARRATIVE_REPLACEMENTS
)

# ── Main ───────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--dry-run", action="store_true",
                        help="Report changes without writing output")
    args = parser.parse_args()

    if not INPUT_DOCX.exists():
        print(f"ERROR: Input not found: {INPUT_DOCX}", file=sys.stderr)
        sys.exit(1)

    if not TABLE7_CSV.exists():
        print(f"ERROR: Table7 CSV not found: {TABLE7_CSV}", file=sys.stderr)
        sys.exit(1)

    # Backup
    if not args.dry_run:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        bak = PAPERS / f"{INPUT_DOCX.stem}.bak_{ts}{INPUT_DOCX.suffix}"
        shutil.copy2(INPUT_DOCX, bak)
        print(f"Backup: {bak.name}")

    doc = Document(str(INPUT_DOCX))

    # --- Text replacements ---
    print("\n=== Text replacements ===")
    counts = apply_replacements(doc, ALL_REPLACEMENTS, dry_run=args.dry_run)
    total_replaced = 0
    for old, count in counts.items():
        if count > 0:
            print(f"  [{count}x] {old[:80]!r}")
            total_replaced += count
        else:
            print(f"  [0x] NOT FOUND: {old[:80]!r}")

    # --- Table 7 update ---
    print("\n=== Table 7 update ===")
    rows_csv = _load_table7_csv()
    logs = update_table7(doc, rows_csv, dry_run=args.dry_run)
    for log in logs:
        print(f"  {log}")

    # --- Write output ---
    print(f"\nTotal text replacements: {total_replaced}")
    if args.dry_run:
        print("\nDRY RUN — no file written.")
    else:
        doc.save(str(OUTPUT_DOCX))
        print(f"\nWritten: {OUTPUT_DOCX.name}")

    # --- Post-check: verify no Portaria 69/2021 remains ---
    if not args.dry_run:
        doc2 = Document(str(OUTPUT_DOCX))
        all_text = []
        for para in doc2.paragraphs:
            all_text.append(para.text)
        for table in doc2.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)
        combined = "\n".join(all_text)
        remaining_69 = combined.count("69/2021")
        remaining_59b = combined.count("59-B")
        print(f"\nPost-check:")
        print(f"  '69/2021' remaining: {remaining_69} (target: 0)")
        print(f"  '59-B' remaining:    {remaining_59b} (target: 0)")
        if remaining_69 > 0 or remaining_59b > 0:
            print("  WARNING: manual review required for remaining instances")
        else:
            print("  OK — all critical strings replaced")


if __name__ == "__main__":
    main()
