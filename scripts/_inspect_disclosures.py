"""Inspect mandatory disclosure area and check for psi_N table in docx."""
from docx import Document

doc = Document("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")
paras = doc.paragraphs

# Show paragraphs 295-370 (around the disclosure area)
print("=== DISCLOSURE AREA (paras 295-370) ===")
for i in range(295, min(370, len(paras))):
    t = paras[i].text.strip()
    s = (paras[i].style.name or "")[:20]
    if t:
        print(f"[{i:4d}] ({s:20}) {t[:120].encode('ascii', 'replace').decode()}")

# Check if psi_N source table already exists (look in tables)
print("\n=== TABLES CHECK ===")
for j, tbl in enumerate(doc.tables):
    rows = tbl.rows
    if rows:
        header_text = " | ".join(c.text for c in rows[0].cells)[:100]
        print(f"Table {j}: {len(rows)} rows x {len(rows[0].cells)} cols | Header: {header_text.encode('ascii','replace').decode()}")
        # Check if it could be the psi_N table
        all_text = " ".join(c.text for row in rows for c in row.cells).lower()
        if "synthetic" in all_text or "psi" in all_text or "source" in all_text:
            print(f"  >> CANDIDATE psi_N table")
            for row in rows:
                print(f"     {' | '.join(c.text[:30] for c in row.cells).encode('ascii','replace').decode()}")
