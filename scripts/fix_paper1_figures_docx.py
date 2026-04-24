#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Cirurgia in-place de figuras no PAPER1_QFENG_FINAL_editando.docx.

Operações:
  1. image1 (F3_hilbert→F1_interference_overview) – blob swap + rewrite caption para[217]
  2. image2 (F2_manaus_theta→F3_hilbert) – blob swap + MOVE de antes-da-Table3 p/ após Fig1
     + inserir novo parágrafo de caption Fig 2
  3. image3 (F2_manaus_timeseries→F2_manaus_theta_efetivo v2) – blob swap + rewrite caption
  4. image4 (F4, mesma imagem) – só rewrite caption para[310]
  5. placeholder [INSERT FIGURE 5 HERE] → imagem F7 + rewrite caption
  6. placeholder [INSERT FIGURE 6 HERE] → imagem F5 + rewrite caption
  7. placeholder [INSERT FIGURE 7 HERE] → imagem F6 + rewrite caption

Usage:
    python scripts/fix_paper1_figures_docx.py \\
        --docx   docs/papers/PAPER1_QFENG_FINAL_editando.docx \\
        --figdir outputs/figures
"""
from __future__ import annotations

import argparse
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.oxml import OxmlElement  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from docx.shared import Inches  # type: ignore[import-untyped]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

# ── Captions finais (literais, aprovadas pelo usuário) ──────────────────────────
CAPTIONS: dict[int, str] = {
    1: (
        "Interference angle θ across seven scenarios — overview of governance regime "
        "classification. Predictor states ψ_N (dashed) plotted by their angular separation "
        "from the normative reference ψ_S (solid). Five CIRCUIT-BREAKER scenarios cluster at "
        "127.8°–134.7° (destructive interference); two positive controls fall within the STAC "
        "band at 5.7°–7.1° (constructive interference)."
    ),
    2: (
        "Q-FENG Interference Geometry in the Decision Hilbert Space. Angle θ between predictor "
        "state ψ_N (dashed) and normative state ψ_S (solid) across seven scenarios and two "
        "normative regimes. Health Governance: C2, C3, C7 (ψ in R³, 3 actions). Labour Law: "
        "T-CLT-01 through T-CLT-04 (ψ in R², 2 actions). GSP annotation shows governance "
        "suppression percentage per scenario."
    ),
    3: (
        "Markovian θ_eff Trajectory — Manaus COVID-19 Health Crisis (Jul 2020 – Jun 2021). "
        "Circuit-Breaker activated October 2020 (θ_eff = 125.3°, α(t) = 0.909), three months "
        "before the January 2021 ICU collapse declared by Portaria MS 69/2021. Left axis: θ_eff "
        "Markovian (SIH/DATASUS) and θ_t instantaneous, with 95% bootstrap CI shading. Right "
        "axis: hospital occupancy rate (%). Peak February 2021 at θ_eff = 130.9°."
    ),
    4: (
        "Governance Suppression Percentage by scenario — Born-rule Quantum vs. Classical "
        "Bayesian. GSP quantifies the suppression of the norm-violating action probability by the "
        "quantum interference formalism, relative to a classical Bayesian mixture model. CB "
        "scenarios (θ ≥ 120°): destructive interference, GSP ∈ [9.4%, 25.2%]. STAC positive "
        "controls (θ < 30°): constructive interference, GSP ∈ [−0.44%, −0.28%]."
    ),
    5: (
        "DeonticAtoms per applied track and modality distribution. Panel (a): atoms extracted per "
        "normative track (Brazil health, EU health, USA health, Brazil labour). Panel (b): "
        "modality distribution (obligation, permission, prohibition, faculty) by track. Total: "
        "10,142 DeonticAtoms at E2 (5,136 health/governance; 5,006 labour)."
    ),
    6: (
        "Threshold Robustness — STAC/CB Classification Stability Across Parameter Grid. Grid "
        "search over θ_stac ∈ {20,25,30,35,40}° and θ_block ∈ {100,105,...,130}° (5×7=35 "
        "combinations per scenario; 245 total evaluations). Overall robustness: 97.96% (240/245). "
        "At θ_block ≤ 125°: 100% stability. Empirical θ gap: [7.0°, 127.8°]."
    ),
    7: (
        "ψ-Weight Sensitivity Analysis — Monte Carlo Robustness Under ±20% Perturbation. For "
        "each scenario, 500 perturbation samples were drawn by adding U(−δ,+δ) noise to each "
        "component of ψ_N (δ=20%), re-normalising, and recomputing θ. Correct regime "
        "preservation: 100% across all 7 scenarios (3500 total samples). Maximum σ_θ: 2.01° "
        "(T-CLT-02)."
    ),
}

# ── Figuras: rId → PNG destino ──────────────────────────────────────────────────
IMAGE_SWAPS: dict[str, str] = {
    "rId9":  "F1_interference_overview.png",   # image1: F3_hilbert → F1
    "rId10": "F3_hilbert_decision_space.png",  # image2: F2_manaus_theta → F3_hilbert
    "rId11": "F2_manaus_theta_efetivo.png",    # image3: F2_timeseries → F2_manaus_theta v2
    # rId12 (image4 = F4_governance) mantém blob; só caption muda
}


def _abort(msg: str) -> None:
    print(f"ABORT: {msg}", file=sys.stderr)
    sys.exit(1)


def _assert_count_1(paragraphs: list, anchor: str) -> int:
    """Return the index of the single paragraph containing anchor, or abort."""
    hits = [i for i, p in enumerate(paragraphs) if anchor in p.text]
    if len(hits) != 1:
        _abort(f"Expected exactly 1 paragraph containing {anchor!r}, found {len(hits)}: {hits}")
    return hits[0]


def _rewrite_caption(para: Paragraph, fig_num: int) -> None:
    """Clear runs, write 'Figure N. ' bold + caption body."""
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    r1 = para.add_run(f"Figure {fig_num}. ")
    r1.bold = True
    para.add_run(CAPTIONS[fig_num])


def _new_paragraph_after(anchor_p: Paragraph, doc: Document, style_name: str = "Caption") -> Paragraph:
    """Insert a new empty paragraph immediately after anchor_p."""
    new_p_elem = OxmlElement("w:p")
    anchor_p._p.addnext(new_p_elem)
    new_para = Paragraph(new_p_elem, doc)
    try:
        new_para.style = doc.styles[style_name]
    except KeyError:
        pass
    return new_para


def _replace_placeholder_with_image(
    para: Paragraph,
    doc: Document,
    fig_path: Path,
    width: Inches = Inches(6.0),
) -> None:
    """Clear para runs (keeping style/alignment) and add an inline picture."""
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(fig_path), width=width)


def _rewrite_normal_caption(para: Paragraph, fig_num: int) -> None:
    """Same as _rewrite_caption but also centers the paragraph."""
    for run in list(para.runs):
        run._r.getparent().remove(run._r)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = para.add_run(f"Figure {fig_num}. ")
    r1.bold = True
    r1.italic = True
    para.add_run(CAPTIONS[fig_num]).italic = True


def run_surgery(docx_path: Path, figdir: Path) -> None:
    # ── Pré-condições ──────────────────────────────────────────────────────────
    if not docx_path.exists():
        _abort(f"{docx_path} not found")
    for fn in IMAGE_SWAPS.values():
        if not (figdir / fn).exists():
            _abort(f"Figure not found: {figdir / fn}")
    for fn in ["F4_governance_suppression.png", "F7_deontic_regime_modality.png",
               "F5_threshold_robustness.png", "F6_psi_sensitivity.png"]:
        if not (figdir / fn).exists():
            _abort(f"Figure not found: {figdir / fn}")

    # ── Backup ─────────────────────────────────────────────────────────────────
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = docx_path.with_name(docx_path.stem + f".bak_{ts}.docx")
    shutil.copy2(docx_path, bak)
    print(f"Backup: {bak}", file=sys.stderr)

    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    # ── Op 1: blob swap image1 (rId9: F3_hilbert → F1_interference_overview) ──
    print("Op 1: blob swap rId9 → F1_interference_overview", file=sys.stderr)
    part1 = doc.part.related_parts["rId9"]
    part1._blob = (figdir / "F1_interference_overview.png").read_bytes()

    # Rewrite caption: para with Caption style before image1 image paragraph
    # Known: para[217] "Figure . Interference angle θ by scenario..."
    idx_cap1 = _assert_count_1(paragraphs, "Figure . Interference angle")
    _rewrite_caption(paragraphs[idx_cap1], 1)
    print(f"  Caption rewritten at para[{idx_cap1}]", file=sys.stderr)

    # ── Op 2: blob swap image2 (rId10: F2_manaus → F3_hilbert) + MOVE ─────────
    # image2 is at para[303] (before Table 3, no caption) — move to right after image1
    print("Op 2: blob swap rId10 → F3_hilbert + MOVE to after Fig1", file=sys.stderr)
    part2 = doc.part.related_parts["rId10"]
    part2._blob = (figdir / "F3_hilbert_decision_space.png").read_bytes()

    # Find image1 paragraph (has blip rId9 in its XML)
    img1_p = None
    for p in paragraphs:
        if 'rId9' in p._p.xml and "a:blip" in p._p.xml:
            img1_p = p
            break
    if img1_p is None:
        _abort("Could not locate image1 paragraph (rId9)")

    # Find image2 paragraph (has blip rId10 in its XML)
    img2_p = None
    for p in paragraphs:
        if "rId10" in p._p.xml and "a:blip" in p._p.xml:
            img2_p = p
            break
    if img2_p is None:
        _abort("Could not locate image2 paragraph (rId10)")

    # Detach image2 from current location
    img2_p._p.getparent().remove(img2_p._p)

    # Insert after image1
    img1_p._p.addnext(img2_p._p)

    # Insert new Fig 2 caption immediately after the moved image2
    # (img2_p._p is now after img1_p._p)
    fig2_caption = _new_paragraph_after(img2_p, doc, "Caption")
    _rewrite_caption(fig2_caption, 2)
    print("  image2 moved after Fig1 and Fig2 caption inserted", file=sys.stderr)

    # ── Op 3: blob swap image3 (rId11: F2_timeseries → F2_manaus_theta_efetivo) ──
    print("Op 3: blob swap rId11 → F2_manaus_theta_efetivo", file=sys.stderr)
    part3 = doc.part.related_parts["rId11"]
    part3._blob = (figdir / "F2_manaus_theta_efetivo.png").read_bytes()

    # Rewrite adjacent caption: Normal-style para with "Figure 3. Manaus"
    # After the move in Op 2, paragraph indices shifted — search by text instead
    hits3 = [p for p in doc.paragraphs if "Figure 3. Manaus" in p.text or
             ("Figure 3" in p.text and "Manaus" in p.text)]
    if not hits3:
        # Also check for any "Figure 3" caption adjacent to image3
        img3_p = None
        for p in doc.paragraphs:
            if "rId11" in p._p.xml and "a:blip" in p._p.xml:
                img3_p = p
                break
        if img3_p:
            paras_list = list(doc.paragraphs)
            idx3 = paras_list.index(img3_p)
            # Check next para
            if idx3 + 1 < len(paras_list) and "Figure" in paras_list[idx3 + 1].text:
                hits3 = [paras_list[idx3 + 1]]
    if hits3:
        _rewrite_normal_caption(hits3[0], 3)
        print(f"  Caption Fig3 rewritten", file=sys.stderr)
    else:
        # Insert new caption after image3
        img3_p2 = None
        for p in doc.paragraphs:
            if "rId11" in p._p.xml and "a:blip" in p._p.xml:
                img3_p2 = p
                break
        if img3_p2:
            new_cap3 = _new_paragraph_after(img3_p2, doc, "Caption")
            _rewrite_caption(new_cap3, 3)
            print("  Inserted new Fig3 caption", file=sys.stderr)

    # ── Op 4: rewrite caption for image4 (rId12, F4_governance — no blob change) ──
    print("Op 4: rewrite caption for image4 (rId12)", file=sys.stderr)
    hits4 = [p for p in doc.paragraphs if "Figure . Governance suppression" in p.text]
    if not hits4:
        hits4 = [p for p in doc.paragraphs if "Governance suppression" in p.text
                 and p.style and "Caption" in p.style.name]
    if not hits4:
        _abort("Could not locate Figure 4 caption paragraph")
    _rewrite_caption(hits4[0], 4)
    print("  Caption Fig4 rewritten", file=sys.stderr)

    # ── Ops 5/6/7: replace placeholders with images ────────────────────────────
    placeholder_ops = [
        ("[INSERT FIGURE 5 HERE]", "F7_deontic_regime_modality.png", "Figure 5. DeonticAtom", 5),
        ("[INSERT FIGURE 6 HERE]", "F5_threshold_robustness.png",    "Figure 6.",              6),
        ("[INSERT FIGURE 7 HERE]", "F6_psi_sensitivity.png",         "Figure 7.",              7),
    ]

    for placeholder_text, fig_filename, old_caption_prefix, fig_num in placeholder_ops:
        print(f"Op {fig_num}: insert {fig_filename} at {placeholder_text!r}", file=sys.stderr)

        # Locate placeholder paragraph by text (fresh scan each time)
        placeholder_elem = None
        for p in doc.paragraphs:
            if p.text.strip() == placeholder_text:
                placeholder_elem = p._p  # store XML element, not Python wrapper
                break
        if placeholder_elem is None:
            _abort(f"Placeholder not found: {placeholder_text!r}")

        # Re-wrap as Paragraph for manipulation
        placeholder_p = Paragraph(placeholder_elem, doc)
        _replace_placeholder_with_image(placeholder_p, doc, figdir / fig_filename)

        # Find adjacent caption using XML next-sibling walk (immune to list re-creation)
        caption_p = None
        sibling = placeholder_elem.getnext()
        for _ in range(4):
            if sibling is None:
                break
            candidate = Paragraph(sibling, doc)
            txt = candidate.text.strip()
            if f"Figure {fig_num}" in txt or old_caption_prefix.rstrip(".") in txt:
                caption_p = candidate
                break
            sibling = sibling.getnext()

        if caption_p is None:
            caption_p = _new_paragraph_after(placeholder_p, doc, "Caption")
        _rewrite_normal_caption(caption_p, fig_num)
        print(f"  Fig{fig_num} inserted and caption rewritten", file=sys.stderr)

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(docx_path)
    print(f"Saved: {docx_path}", file=sys.stderr)

    # ── Post-verification ─────────────────────────────────────────────────────
    print("\n=== POST-VERIFICATION ===", file=sys.stderr)
    doc2 = Document(docx_path)
    file_size = docx_path.stat().st_size
    bak_size = bak.stat().st_size
    print(f"File size: {file_size:,} bytes (backup: {bak_size:,})", file=sys.stderr)
    if file_size <= bak_size:
        print("WARNING: file shrank — verify manually!", file=sys.stderr)

    # Count images
    n_images = sum(1 for p in doc2.paragraphs if "a:blip" in p._p.xml)
    print(f"Images (InlineShape blip refs): {n_images}", file=sys.stderr)

    # Count remaining placeholders
    placeholders_remaining = [
        p.text for p in doc2.paragraphs if "INSERT FIGURE" in p.text
    ]
    if placeholders_remaining:
        print(f"WARNING: Remaining placeholders: {placeholders_remaining}", file=sys.stderr)
    else:
        print("OK: No remaining [INSERT FIGURE] placeholders", file=sys.stderr)

    # Check captions
    all_ok = True
    for fig_num, caption_body in CAPTIONS.items():
        expected = f"Figure {fig_num}. " + caption_body[:40]
        found = any(expected[:40] in p.text for p in doc2.paragraphs)
        status = "OK" if found else "MISSING"
        if not found:
            all_ok = False
        print(f"  Caption {fig_num}: {status}", file=sys.stderr)

    # Check Herrera / Díaz-Rodríguez still present
    text_all = " ".join(p.text for p in doc2.paragraphs)
    for name in ["Herrera", "Díaz-Rodríguez"]:
        if name in text_all:
            print(f"OK: {name!r} still present", file=sys.stderr)
        else:
            print(f"WARNING: {name!r} NOT FOUND — check regressão de literatura!", file=sys.stderr)
            all_ok = False

    if all_ok:
        print("\n✓ All post-checks passed.", file=sys.stderr)
    else:
        print("\n✗ Some checks failed — review output.", file=sys.stderr)
        sys.exit(2)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fix figure images/captions in PAPER1 docx")
    parser.add_argument("--docx",   required=True, help="Path to .docx (edited in place)")
    parser.add_argument("--figdir", required=True, help="Directory containing F1..F7 PNG files")
    args = parser.parse_args()

    run_surgery(Path(args.docx), Path(args.figdir))


if __name__ == "__main__":
    main()
