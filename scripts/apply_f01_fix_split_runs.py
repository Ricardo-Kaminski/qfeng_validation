"""Fix split-run occurrences of fabricated TST case number in the auditfixed docx.

Handles the case where TST-RR-000200-50.2019.5.02.0020 is spread across multiple
<w:t> nodes within a single <w:p> paragraph.
"""
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import etree

DOCX_PATH = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")

OLD_CASE = "TST-RR-000200-50.2019.5.02.0020"
NEW_CASE = "TST-Ag-RR-868-65.2021.5.13.0030"

OLD_SHORT = "RR-000200-50.2019"
NEW_SHORT = "Ag-RR-868-65.2021"

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _get_para_text(para_elem) -> str:
    return "".join(t.text or "" for t in para_elem.iter(f"{{{WNS}}}t"))


def fix_split_runs_in_para(para_elem, old: str, new: str) -> int:
    """Replace old→new across merged text of a paragraph, consolidating runs."""
    full_text = _get_para_text(para_elem)
    if old not in full_text:
        return 0

    replaced_text = full_text.replace(old, new)
    count = full_text.count(old)

    # Get all w:t nodes in this paragraph
    t_nodes = list(para_elem.iter(f"{{{WNS}}}t"))
    if not t_nodes:
        return 0

    # Put entire replaced text in first w:t, clear the rest
    # (preserves the first run's formatting for the whole text)
    t_nodes[0].text = replaced_text
    # Ensure xml:space="preserve" so leading/trailing spaces are kept
    t_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    for t in t_nodes[1:]:
        t.text = ""

    return count


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"Not found: {DOCX_PATH}")

    bak = DOCX_PATH.with_suffix(f".bak_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    shutil.copy2(DOCX_PATH, bak)
    print(f"Backup: {bak}")

    doc = Document(DOCX_PATH)
    total = 0

    # Iterate all paragraphs including those inside tables
    from docx.oxml.ns import qn
    all_paras = doc.element.body.iter(f"{{{WNS}}}p")
    for para_elem in all_paras:
        para_text = _get_para_text(para_elem)
        if OLD_CASE in para_text:
            n = fix_split_runs_in_para(para_elem, OLD_CASE, NEW_CASE)
            if n:
                print(f"  [OK] Split-run fix: '{OLD_CASE}' → '{NEW_CASE}' ({n}×)")
                print(f"       Context: '{para_text[:80].strip()}'")
                total += n
        elif OLD_SHORT in para_text:
            n = fix_split_runs_in_para(para_elem, OLD_SHORT, NEW_SHORT)
            if n:
                print(f"  [OK] Split-run fix: '{OLD_SHORT}' → '{NEW_SHORT}' ({n}×)")
                total += n

    doc.save(DOCX_PATH)
    print(f"\nSaved: {DOCX_PATH}")
    print(f"Total split-run fixes: {total}")

    # Post-check
    doc2 = Document(DOCX_PATH)
    remaining_text = ""
    for p in doc2.paragraphs:
        remaining_text += "".join(r.text for r in p.runs)
    for tbl in doc2.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    remaining_text += "".join(r.text for r in p.runs)

    for check_str in ["000200-50.2019", "RR-000200", "tst_rr_000200"]:
        occ = remaining_text.count(check_str)
        status = "[OK]" if occ == 0 else "[WARN]"
        print(f"  {status} Post-check: '{check_str}' → {occ} occurrences")


if __name__ == "__main__":
    main()
