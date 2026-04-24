#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract PAPER1_QFENG_FINAL_editando.docx to a Markdown mirror.

Usage:
    python scripts/extract_docx_to_md.py \\
        --input  docs/papers/PAPER1_QFENG_FINAL_editando.docx \\
        --output docs/papers/PAPER1_QFENG_FINAL_editando.md
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from lxml import etree  # type: ignore[import-untyped]

HEADING_MAP = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Title": "#",
}


def _para_has_image(p_elem: etree._Element) -> bool:
    # String-based check is reliable across namespace prefix variations
    xml_str: str = etree.tostring(p_elem, encoding="unicode")
    return "a:blip" in xml_str


def _get_image_marker(p_elem: etree._Element, doc: Document) -> str:
    import re as _re
    xml_str: str = etree.tostring(p_elem, encoding="unicode")
    # Extract r:embed="rIdN" from blip elements
    rids = _re.findall(r'r:embed="(rId\d+)"', xml_str)
    if not rids:
        rids = _re.findall(r'r:link="(rId\d+)"', xml_str)
    markers = []
    for rid in rids:
        if rid in doc.part.related_parts:
            part = doc.part.related_parts[rid]
            partname = str(part.partname)
            size = len(part.blob)
            markers.append(
                f"![Figure]({partname})"
                f"{{#fig-{rid} rId={rid} size={size:,}b}}"
            )
        else:
            markers.append(f"![Figure](unknown){{#fig-{rid}}}")
    return "\n".join(markers)


def _table_to_md(table) -> str:  # type: ignore[no-untyped-def]
    rows = table.rows
    if not rows:
        return ""
    lines = []
    for r_idx, row in enumerate(rows):
        cells = [c.text.replace("|", "\\|").replace("\n", " ").strip() for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if r_idx == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(lines)


def extract_to_md(docx_path: Path, md_path: Path) -> None:
    doc = Document(docx_path)

    # Build a list of body children in document order (paragraphs + tables)
    body = doc.element.body
    para_objs = {id(p._p): p for p in doc.paragraphs}
    table_objs = {id(t._tbl): t for t in doc.tables}

    chunks: list[str] = []

    for child in body:
        tag = etree.QName(child.tag).localname

        if tag == "p":
            p_obj = para_objs.get(id(child))
            if p_obj is None:
                continue

            style_name = p_obj.style.name if p_obj.style else "Normal"
            text = p_obj.text or ""

            if _para_has_image(child):
                chunks.append(_get_image_marker(child, doc))
                continue

            if not text.strip():
                chunks.append("")
                continue

            if style_name in HEADING_MAP:
                prefix = HEADING_MAP[style_name]
                chunks.append(f"{prefix} {text}")
            elif style_name == "Caption":
                chunks.append(f"*{text}*")
            else:
                chunks.append(text)

        elif tag == "tbl":
            t_obj = table_objs.get(id(child))
            if t_obj is None:
                continue
            chunks.append(_table_to_md(t_obj))

    md_content = "\n\n".join(c for c in chunks if c is not None)
    md_path.parent.mkdir(parents=True, exist_ok=True)
    md_path.write_text(md_content, encoding="utf-8")
    print(f"Extracted {len(doc.paragraphs)} paragraphs, "
          f"{len(doc.tables)} tables → {md_path}", file=sys.stderr)


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract docx → markdown mirror")
    parser.add_argument("--input", required=True, help="Path to .docx")
    parser.add_argument("--output", required=True, help="Path to output .md")
    args = parser.parse_args()

    docx_path = Path(args.input)
    md_path = Path(args.output)

    if not docx_path.exists():
        print(f"ERROR: {docx_path} not found", file=sys.stderr)
        sys.exit(1)

    extract_to_md(docx_path, md_path)
    print(f"OK — written to {md_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
