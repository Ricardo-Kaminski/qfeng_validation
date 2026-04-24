#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Quick diagnostic — inspect internal structure of PAPER1_QFENG_FINAL_editando.docx."""
import io
import re
import sys
import zipfile
from pathlib import Path

# Force UTF-8 output on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

DOCX = Path(r"C:\Workspace\academico\qfeng_validacao\docs\papers\PAPER1_QFENG_FINAL_editando.docx")


def main() -> None:
    with zipfile.ZipFile(DOCX) as z:
        # 1. Media files
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        print("=== MEDIA FILES ===")
        for m in sorted(media):
            info = z.getinfo(m)
            print(f"  {m}  ({info.file_size:,} bytes)")

        # 2. Relationships
        print("\n=== IMAGE RELATIONSHIPS ===")
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")
        image_rels: dict[str, str] = {}
        # Try multiple attribute orderings
        for pat in [
            r'Id="(rId\d+)"[^>]*Target="([^"]+)"[^>]*Type="([^"]+)"',
            r'Type="([^"]+)"[^>]*Id="([^"]+)"[^>]*Target="([^"]+)"',
        ]:
            for m in re.finditer(pat, rels_xml):
                g = m.groups()
                if "image" in g[2].lower() or "media" in g[1].lower():
                    print(f"  {g[0]} -> {g[1]} [{g[2].split('/')[-1]}]")
                    image_rels[g[0]] = g[1]
        # Fallback: just grep for image lines
        if not image_rels:
            for line in rels_xml.split(">"):
                if "media" in line.lower() or ("image" in line.lower() and "Id=" in line):
                    print(f"  RAW: {line.strip()[:200]}")

        # 3. Scan document.xml for drawings + adjacent paragraphs
        print("\n=== PARAGRAPH/DRAWING STRUCTURE (around images) ===")
        doc_xml = z.read("word/document.xml").decode("utf-8")

        # Find all blip embeds (inline image references)
        for m in re.finditer(r'r:embed="(rId\d+)"', doc_xml):
            print(f"  blip embed: {m.group(1)}")

        # 4. Find paragraphs with placeholder text
        print("\n=== PLACEHOLDER PARAGRAPHS ===")
        from docx import Document  # type: ignore
        doc = Document(DOCX)
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if "INSERT FIGURE" in txt or "Figure" in txt or "figure" in txt:
                style = p.style.name if p.style else "?"
                print(f"  para[{i:4d}] style={style!r:25s} | {txt[:120]!r}")

        # 5. Show styles used for captions
        print("\n=== CAPTION-STYLE PARAGRAPHS ===")
        for i, p in enumerate(doc.paragraphs):
            if p.style and ("caption" in p.style.name.lower() or "Caption" in p.style.name):
                print(f"  para[{i:4d}] style={p.style.name!r} | {p.text[:100]!r}")

        # 6. Find image para indices and their neighbors
        print("\n=== IMAGE PARAGRAPH LOCATIONS ===")
        from docx.oxml.ns import qn  # type: ignore
        for i, p in enumerate(doc.paragraphs):
            has_drawing = p._p.find(qn("w:r")) is not None and (
                p._p.find(f".//{{{qn('a:blip')}}}" ) is not None  # noqa
                or b"<a:blip" in p._p.xml.encode() if hasattr(p._p, "xml") else False
            )
            # simpler check
            if "a:blip" in p._p.xml:
                style = p.style.name if p.style else "?"
                print(f"  para[{i:4d}] HAS IMAGE  style={style!r}")
                # print neighbors
                for delta, label in [(-2, "prev-2"), (-1, "prev-1"), (1, "next+1"), (2, "next+2")]:
                    j = i + delta
                    if 0 <= j < len(doc.paragraphs):
                        np_ = doc.paragraphs[j]
                        nstyle = np_.style.name if np_.style else "?"
                        has_img = "a:blip" in np_._p.xml
                        txt = np_.text[:80]
                        print(f"    {label:8s} para[{j:4d}] style={nstyle!r:25s} img={has_img} | {txt!r}")

        # 7. Anchors
        print("\n=== KEY TEXT ANCHORS ===")
        anchors = [
            "Overview of Seven Scenarios",
            "Table 2",
            "Manaus Theta",
            "Table 3",
            "Born-Rule",
            "GSP",
            "Table 4",
            "Table 6",
            "Table 7",
            "INSERT FIGURE",
            "Herrera",
            "Diaz-Rodriguez",
            "Díaz-Rodríguez",
        ]
        for i, p in enumerate(doc.paragraphs):
            for anchor in anchors:
                if anchor.lower() in p.text.lower():
                    style = p.style.name if p.style else "?"
                    print(f"  [{anchor!r:35s}] para[{i:4d}] style={style!r:25s} | {p.text[:80]!r}")
                    break


if __name__ == "__main__":
    main()
