"""Apply ARS HIGH-priority corrections to canonical paper DOCX.

ARS-H1: C7 narrative — replace 'absent from statute' with 'not operationalized'
ARS-H2: T-CLT-02 — reclassify failure_type constitutional → execution_absent_channel
ARS-H3: Z derivation — add explicit derivation sentence at Para 123
"""
from __future__ import annotations

import datetime
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")


def backup(path: Path) -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dst = path.with_suffix(f".bak_{ts}.docx")
    shutil.copy2(path, dst)
    print(f"Backup: {dst}")
    return dst


def para_idx(doc: docx.Document, fragment: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    raise ValueError(f"Fragment not found: {fragment!r}")


def replace_para_text(doc: docx.Document, idx: int, new_text: str) -> None:
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"  Para {idx} replaced.")


def insert_paragraph_after(doc: docx.Document, ref_idx: int,
                            text: str, style_id: str = "Normal") -> None:
    ref_para = doc.paragraphs[ref_idx]
    new_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    new_para.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    ref_para._element.addnext(new_para)


# ══════════════════════════════════════════════════════════════════════
# ARS-H1: C7 — 'absent from statute' → 'not operationalized'
# ══════════════════════════════════════════════════════════════════════

C7_FAILURE_OLD = (
    "\tThe failure type is constitutional: 14th Amendment §1 EPC + 42 U.S.C. §2000d requires "
    "that Medicaid services be provided “consistent with the best interests of the beneficiaries” "
    "but does not specify a racial equity obligation by name. The equal-protection clause enters "
    "the normative architecture exclusively through the 14th Amendment grounding — a constitutional "
    "interpretive gap that the commercial algorithm structurally exploits. The corrective action is "
    "not statutory amendment but judicial interpretation: the equal-protection constraint must be "
    "formally incorporated into the Medicaid state plan requirements through either Congressional "
    "action or court ruling. This is a constitutional failure in the precise sense: the statute does "
    "not contain the sovereign predicate that constitutional law requires it to contain."
)

C7_FAILURE_NEW = (
    "\tThe failure type is constitutional: the equal-protection obligation encoded in the 14th "
    "Amendment §1 EPC and Title VI §601 (42 U.S.C. §2000d) is present in the normative corpus "
    "as active sovereign predicates (equal_protection_of_the_laws, prohibition_disparate_impact_in_federal_programs), "
    "but was never operationalized in the commercial algorithm's governance layer. "
    "The algorithm's System 5 (autonomous decision-making) was designed without a mechanism to "
    "check whether its risk-score assignments produced racially disparate outcomes — the sovereign "
    "predicate existed in constitutional architecture but was not instantiated as a runtime constraint "
    "on the algorithm's output. The corrective action is prospective: the equal-protection constraint "
    "must be inscribed in the algorithm's design and auditing pipeline before deployment, not "
    "retrofitted post hoc through statutory amendment. This is a constitutional failure in the Q-FENG "
    "sense: the sovereign predicate is derivable from the normative corpus but was absent from the "
    "algorithm's pre-deployment design."
)

# ══════════════════════════════════════════════════════════════════════
# ARS-H2: T-CLT-02 — reclassify constitutional → execution_absent_channel
# ══════════════════════════════════════════════════════════════════════

TCLT02_HEADING_OLD = (
    "T-CLT-02 — Hour Bank Without CCT (Brasil, constitutional): "
    "θ = 127.81°, CIRCUIT_BREAKER — the lowest CB theta in the dataset."
)

TCLT02_HEADING_NEW = (
    "T-CLT-02 — Hour Bank Without CCT (Brasil, execution_absent_channel): "
    "θ = 127.81°, CIRCUIT_BREAKER — the lowest CB theta in the dataset."
)

TCLT02_BODY_OLD = (
    "This reflects a scenario with partial normative support: unlike T-CLT-01 (where the predicate "
    "literally does not exist) or C3 (where the constitutional grounding is unambiguous), T-CLT-02 "
    "involves a normative structure that permits the action under certain conditions (with a CCT) but "
    "prohibits it under the present conditions (without a CCT). The CLT framework thus provides a "
    "partial anchor for the predictor — explaining the lower θ — while the missing CCT condition "
    "pushes the scenario firmly into CIRCUIT_BREAKER."
)

TCLT02_BODY_NEW = (
    "This reflects an execution_absent_channel failure: the sovereign predicates governing hour banks "
    "are fully present in the normative corpus (CLT Art. 59 §§2 e 5, Art. 611-A I, TST Súmula 85 V — "
    "all active in the Clingo answer set), but the required execution channel — a Convenção Coletiva de "
    "Trabalho (CCT) or Acordo Coletivo de Trabalho (ACT) registered and filed — was structurally absent "
    "from the contractual configuration. Unlike T-CLT-01 (execution_inertia: predicate exists, agent "
    "failed to act on it) and C3 (constitutional: predicate absent from System 5 of the allocation "
    "model), T-CLT-02 involves a normative permission structure: the CCT channel exists and is "
    "well-defined, but the employer's choice not to activate it left the sovereign constraint "
    "unsatisfied. The CLT framework provides a partial anchor for the predictor — explaining the "
    "lower θ compared to C3/C7 — while the absent CCT channel pushes the scenario firmly into "
    "CIRCUIT_BREAKER."
)

# ══════════════════════════════════════════════════════════════════════
# ARS-H3: Z derivation — add explicit sentence at Para 123
# ══════════════════════════════════════════════════════════════════════

Z_PARA_FRAGMENT = "where the normalisation factor Z incorporates the quantum interference cross-term:"

Z_PARA_OLD = "where the normalisation factor Z incorporates the quantum interference cross-term:"

Z_PARA_NEW = (
    "where the normalisation factor Z incorporates the quantum interference cross-term. "
    "Z arises from the probability normalisation condition: "
    "Σ_j |αψ_N[j] + βψ_S[j]|² = "
    "α²‖ψ_N‖² + β²‖ψ_S‖² + "
    "2αβ⟨ψ_N|ψ_S⟩ = 1 + 2αβcos(θ), "
    "where the last equality uses ‖ψ_N‖ = ‖ψ_S‖ = 1 (L2-normalised) "
    "and the inner-product definition ⟨ψ_N|ψ_S⟩ = cos(θ):"
)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    backup(DOCX)
    doc = docx.Document(str(DOCX))

    # ── ARS-H1: C7 narrative ─────────────────────────────────────────
    print("\n[ARS-H1] C7 'absent from statute' → 'not operationalized'...")
    idx_c7 = para_idx(doc, "The failure type is constitutional: 14th Amendment")
    print(f"  Found at para {idx_c7}")
    replace_para_text(doc, idx_c7, C7_FAILURE_NEW)

    # ── ARS-H2: T-CLT-02 reclassification ────────────────────────────
    print("\n[ARS-H2] T-CLT-02 reclassification constitutional → execution_absent_channel...")
    idx_h2 = para_idx(doc, "T-CLT-02 — Hour Bank Without CCT (Brasil, constitutional)")
    print(f"  Heading at para {idx_h2}")
    # Rebuild heading text (heading may have multiple runs — clear all)
    p = doc.paragraphs[idx_h2]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = TCLT02_HEADING_NEW
    else:
        p.add_run(TCLT02_HEADING_NEW)
    print(f"  Heading replaced.")

    idx_h2_body = para_idx(doc, "This reflects a scenario with partial normative support")
    print(f"  Body at para {idx_h2_body}")
    replace_para_text(doc, idx_h2_body, TCLT02_BODY_NEW)

    # ── ARS-H3: Z derivation ─────────────────────────────────────────
    print("\n[ARS-H3] Adding Z derivation...")
    idx_z = para_idx(doc, Z_PARA_FRAGMENT)
    print(f"  Z para at {idx_z}")
    replace_para_text(doc, idx_z, Z_PARA_NEW)

    doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}")

    # Post-check
    doc2 = docx.Document(str(DOCX))
    checks = [
        ("never operationalized in the commercial algorithm", "C7 not-operationalized framing"),
        ("execution_absent_channel", "T-CLT-02 heading reclassification"),
        ("execution_absent_channel failure: the sovereign predicates", "T-CLT-02 body reclassification"),
        ("probability normalisation condition", "Z derivation"),
        ("inner-product definition", "Z derivation cos(θ)"),
    ]
    all_text = "\n".join(p.text for p in doc2.paragraphs)
    print("\nPost-check:")
    ok = True
    for frag, label in checks:
        found = frag in all_text
        status = "OK" if found else "MISSING"
        print(f"  {status}: {label}")
        if not found:
            ok = False
    assert ok, "Post-check failed."
    print("\nAll checks passed.")


if __name__ == "__main__":
    main()
