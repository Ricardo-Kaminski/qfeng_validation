"""
PAPER1 Chicago Author-Date DOCX Generator
Target audience: IEEE, AI governance journals, applied AI
Converts PAPER1_QFENG_VALIDATION.md → PAPER1_QFENG_VALIDATION_CHICAGO_FINAL.docx

Chicago Author-Date (technical variant):
- Times New Roman 12pt, double-spaced, 1-inch margins
- Numbered headings (1. / 1.1 / 1.1.1)
- Page numbers top-right, no running head
- References: hanging indent, Author Year format
- Figures: placeholder blocks, user inserts manually
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE     = Path(__file__).parent.parent
MD_PATH  = BASE / "docs" / "papers" / "PAPER1_QFENG_VALIDATION.md"
OUT_PATH = BASE / "docs" / "papers" / "PAPER1_QFENG_VALIDATION_CHICAGO_FINAL.docx"

# ── Constants ─────────────────────────────────────────────────────────────────
FONT        = "Times New Roman"
SZ          = Pt(12)
SZ_SM       = Pt(10)
MARGIN      = Inches(1)
INDENT_BODY = Inches(0.5)
INDENT_HANG = Inches(0.5)

AUTHOR           = "Ricardo Kaminski, Ph.D."
AFFILIATION      = "Independent Researcher, Brasília, Brazil"
SUBTITLE         = "Empirical Validation Across Brazilian Health, European AI Law, and US Medicaid Regimes"
PAPER_DATE       = "2026"
KEYWORDS         = ("neurosymbolic AI; deontic logic; AI governance; "
                    "quantum decision theory; legal NLP; Clingo; Answer Set Programming")


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document() -> Document:
    doc = Document()
    # Margins
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = MARGIN
    # Default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = SZ
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)
    _add_page_numbers(doc)
    return doc


def _add_page_numbers(doc: Document):
    """Right-aligned page number in header."""
    for sec in doc.sections:
        hdr = sec.header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        p.clear()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.font.name = FONT
        run.font.size = SZ
        for tag, text in [("begin", ""), ("", "PAGE"), ("end", "")]:
            if tag in ("begin", "end"):
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), tag)
                run._r.append(fc)
            else:
                instr = OxmlElement("w:instrText")
                instr.text = text
                run._r.append(instr)


# ── Paragraph helpers ─────────────────────────────────────────────────────────

def _fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None,
         left_indent=None, space_before=Pt(0), space_after=Pt(0),
         double=True):
    f = para.paragraph_format
    f.alignment       = align
    f.space_before    = space_before
    f.space_after     = space_after
    f.line_spacing_rule = WD_LINE_SPACING.DOUBLE if double else WD_LINE_SPACING.SINGLE
    if first_indent is not None:
        f.first_line_indent = first_indent
    if left_indent is not None:
        f.left_indent = left_indent


def _run(para, text, bold=False, italic=False, font=None, size=None,
         color=None) -> object:
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = font or FONT
    r.font.size = size or SZ
    if color:
        r.font.color.rgb = color
    return r


def blank(doc: Document):
    p = doc.add_paragraph()
    _fmt(p)


def page_break(doc: Document):
    doc.add_page_break()


# ── Inline markdown → runs ────────────────────────────────────────────────────

_INLINE = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')

def inline(para, text: str):
    """Render **bold**, *italic*, ***bold-italic***, `code` into runs."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _run(para, text[pos:m.start()])
        if m.group(2):   # ***bold-italic***
            _run(para, m.group(2), bold=True, italic=True)
        elif m.group(3): # **bold**
            _run(para, m.group(3), bold=True)
        elif m.group(4): # *italic*
            _run(para, m.group(4), italic=True)
        elif m.group(5): # `code`
            _run(para, m.group(5), font="Courier New", size=SZ_SM)
        pos = m.end()
    if pos < len(text):
        _run(para, text[pos:])


# ── Title page ────────────────────────────────────────────────────────────────

def title_page(doc: Document, title: str):
    for _ in range(5):
        blank(doc)

    p = doc.add_paragraph()
    _run(p, title, bold=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0))

    if SUBTITLE:
        p = doc.add_paragraph()
        _run(p, SUBTITLE, italic=True)
        _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0))

    blank(doc)

    p = doc.add_paragraph()
    _run(p, AUTHOR)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    _run(p, AFFILIATION, italic=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    _run(p, PAPER_DATE)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)


# ── Abstract page ─────────────────────────────────────────────────────────────

def abstract_page(doc: Document, lines: list[str], kw: str = ""):
    p = doc.add_paragraph()
    _run(p, "Abstract", bold=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Body — no first-line indent (Chicago abstract style)
    body = " ".join(l.strip() for l in lines if l.strip())
    if body:
        p = doc.add_paragraph()
        inline(p, body)
        _fmt(p, first_indent=Pt(0))

    # Keywords
    p = doc.add_paragraph()
    _run(p, "Keywords: ", italic=True)
    _run(p, kw or KEYWORDS)
    _fmt(p, first_indent=INDENT_BODY)

    page_break(doc)


# ── Headings ──────────────────────────────────────────────────────────────────

def h1(doc: Document, text: str):
    """Chicago Level 1: bold, left-aligned, numbered (e.g. '1. Introduction')."""
    p = doc.add_paragraph()
    _run(p, text, bold=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=Pt(18), space_after=Pt(0))


def h2(doc: Document, text: str):
    """Chicago Level 2: bold, left-aligned (e.g. '1.1 Background')."""
    p = doc.add_paragraph()
    _run(p, text, bold=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=Pt(12), space_after=Pt(0))


def h3(doc: Document, text: str):
    """Chicago Level 3: bold italic, left-aligned (e.g. '1.1.1 Details')."""
    p = doc.add_paragraph()
    _run(p, text, bold=True, italic=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=Pt(6), space_after=Pt(0))


def h4(doc: Document, text: str):
    """Chicago Level 4: italic, left-aligned."""
    p = doc.add_paragraph()
    _run(p, text, italic=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(6))


# ── Body paragraph ────────────────────────────────────────────────────────────

def body(doc: Document, text: str, indent: bool = True):
    p = doc.add_paragraph()
    inline(p, text)
    _fmt(p, first_indent=INDENT_BODY if indent else Pt(0))


# ── Reference entry ───────────────────────────────────────────────────────────

def reference(doc: Document, text: str):
    p = doc.add_paragraph()
    inline(p, text)
    _fmt(p, first_indent=-INDENT_HANG, left_indent=INDENT_HANG)


# ── List items ────────────────────────────────────────────────────────────────

def list_item(doc: Document, text: str, num: int = 0):
    p  = doc.add_paragraph()
    lbl = f"{num}." if num else "•"
    _run(p, f"{lbl}\t")
    inline(p, text)
    _fmt(p, left_indent=INDENT_BODY,
         first_indent=-Inches(0.25))


# ── Code block ────────────────────────────────────────────────────────────────

def code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        _run(p, line or " ", font="Courier New", size=SZ_SM)
        _fmt(p, left_indent=INDENT_BODY, double=False)


# ── Figure placeholder ────────────────────────────────────────────────────────

_GREY = RGBColor(0x88, 0x88, 0x88)
_RED  = RGBColor(0xC0, 0x00, 0x00)

def figure_placeholder(doc: Document, num: str, caption: str):
    # Top border line
    p = doc.add_paragraph()
    _run(p, "─" * 55, color=_GREY)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))

    # [INSERT FIGURE N HERE]
    p = doc.add_paragraph()
    _run(p, f"[INSERT FIGURE {num} HERE]", bold=True, color=_RED)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER)

    # "Figure N" label + caption
    p = doc.add_paragraph()
    _run(p, f"Figure {num}. ", bold=True)
    if caption:
        _run(p, caption, italic=True)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
         first_indent=Pt(0), space_after=Pt(6))

    # Bottom border
    p = doc.add_paragraph()
    _run(p, "─" * 55, color=_GREY)
    _fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


# ── Table ─────────────────────────────────────────────────────────────────────

def md_table(doc: Document, raw_lines: list[str]):
    rows = [l for l in raw_lines
            if not re.match(r'^\|[\s\-:]+\|', l) and l.strip()]
    if not rows:
        return
    parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    n_cols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=n_cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row[:n_cols]):
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            inline(p, cell_text)
            for run in p.runs:
                run.font.name = FONT
                run.font.size = SZ
                run.bold = ri == 0
    blank(doc)


# ── Main parser ───────────────────────────────────────────────────────────────

def parse(doc: Document, md: str):
    lines = md.split("\n")
    n = len(lines)
    i = 0

    title         = ""
    in_abstract   = False
    abstract_buf: list[str] = []
    abstract_done = False
    in_refs       = False
    in_code       = False
    code_buf: list[str] = []
    list_num      = 0

    while i < n:
        line = lines[i]

        # ── HTML comments (skip) ──────────────────────────────────────────────
        if re.match(r'^\s*<!--.*-->\s*$', line):
            i += 1
            continue

        # ── Code fence ────────────────────────────────────────────────────────
        if re.match(r'^(```|~~~)', line):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                code_block(doc, code_buf)
                in_code  = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        m1 = re.match(r'^# (.+)', line)
        m2 = re.match(r'^## (.+)', line)
        m3 = re.match(r'^### (.+)', line)
        m4 = re.match(r'^#### (.+)', line)
        m5 = re.match(r'^##### (.+)', line)

        if m1:
            title = m1.group(1).strip()
            title_page(doc, title)
            i += 1
            continue

        if m2:
            txt = m2.group(1).strip()
            list_num = 0
            if txt.lower() == "abstract":
                in_abstract  = True
                abstract_buf = []
                i += 1
                continue
            # Flush abstract if open
            if in_abstract and not abstract_done:
                abstract_page(doc, abstract_buf)
                abstract_done = True
                in_abstract   = False
            if txt.lower() == "references":
                in_refs = True
                page_break(doc)
                h1(doc, "References")
                i += 1
                continue
            if txt.lower().startswith("appendix"):
                page_break(doc)
            h1(doc, txt)
            i += 1
            continue

        if m3:
            txt = m3.group(1).strip()
            list_num = 0
            if in_abstract and not abstract_done:
                abstract_page(doc, abstract_buf)
                abstract_done = True
                in_abstract   = False
            h2(doc, txt)
            i += 1
            continue

        if m4:
            list_num = 0
            h3(doc, m4.group(1).strip())
            i += 1
            continue

        if m5:
            list_num = 0
            h4(doc, m5.group(1).strip())
            i += 1
            continue

        # ── Figure reference ──────────────────────────────────────────────────
        fig = re.match(r'!\[Figure (\d+):\s*(.*?)\]\(.*?\)', line)
        if fig:
            figure_placeholder(doc, fig.group(1), fig.group(2).strip())
            i += 1
            continue
        img = re.match(r'!\[(.+?)\]\(.*?\)', line)
        if img:
            figure_placeholder(doc, "?", img.group(1))
            i += 1
            continue

        # ── Abstract accumulation ─────────────────────────────────────────────
        if in_abstract:
            abstract_buf.append(line)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        bq = re.match(r'^>\s+(.*)', line)
        if bq:
            p = doc.add_paragraph()
            inline(p, bq.group(1))
            _fmt(p, left_indent=INDENT_BODY, first_indent=Pt(0))
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r'^(-{3,}|\*{3,}|_{3,})\s*$', line):
            blank(doc)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            md_table(doc, tbl_lines)
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        bul = re.match(r'^[-*+]\s+(.+)', line)
        if bul:
            list_item(doc, bul.group(1))
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        num = re.match(r'^\d+\.\s+(.+)', line)
        if num:
            list_num += 1
            list_item(doc, num.group(1), num=list_num)
            i += 1
            continue
        else:
            if not line.strip():
                list_num = 0

        # ── Empty line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── References section ────────────────────────────────────────────────
        if in_refs:
            reference(doc, line.strip())
            i += 1
            continue

        # ── Regular body paragraph ────────────────────────────────────────────
        body(doc, line.strip())
        i += 1


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    md = MD_PATH.read_text(encoding="utf-8")
    doc = setup_document()
    parse(doc, md)
    doc.save(OUT_PATH)
    print(f"Saved  -> {OUT_PATH}")
    print(f"Figures -> search for [INSERT FIGURE N HERE] in red to place images manually.")


if __name__ == "__main__":
    main()
