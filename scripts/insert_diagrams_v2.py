#!/usr/bin/env python3
"""
insert_diagrams_v2.py
Insert 5 new diagrams (4, 6, 7, 8, 9) into Q-FENG paper DOCX copy _diagrams_v1.docx.
Also renumbers existing Diagram 4->5 and Diagram 5->10 in captions and prose.
"""
import io, re, shutil, sys
from pathlib import Path

import cairosvg
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

ROOT  = Path(r"C:\Workspace\academico\qfeng_validacao")
CANON = ROOT / "docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx"
V1    = ROOT / "docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed_diagrams_v1.docx"
FIGS  = ROOT / "artefatos/figuras_paper1"
PROSE = ROOT / "artefatos/briefings/DIAGRAM_INSERTION_PROSE.md"

DPI       = 300
IMG_WIDTH = Inches(5.3)


# ───────────────────────────── helpers ──────────────────────────────


def svg_png(path: Path) -> bytes:
    data = cairosvg.svg2png(url=str(path), dpi=DPI)
    print(f"  {path.name}: {len(data):,} bytes")
    return data


def parse_prose(md: Path) -> dict:
    text = md.read_text(encoding="utf-8")
    out: dict = {}
    key, lines = None, []
    for line in text.splitlines():
        if line.startswith("## INS-"):
            if key and lines:
                out[key] = " ".join(l.strip() for l in lines if l.strip())
            key = line.split()[1]
            lines = []
        elif line.startswith("---"):
            if key and lines:
                out[key] = " ".join(l.strip() for l in lines if l.strip())
            key, lines = None, []
        elif key and line.strip() and not line.startswith("#"):
            lines.append(line)
    if key and lines:
        out[key] = " ".join(l.strip() for l in lines if l.strip())
    return out


def find_elem(doc, fragment: str):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p._element
    raise ValueError(f"Anchor not found: {fragment!r}")


def make_text_elem(doc, text: str, style: str = "Normal"):
    style_ids = {"Normal": "Normal", "Caption": "Legenda"}
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pSt = OxmlElement("w:pStyle")
    pSt.set(qn("w:val"), style_ids.get(style, style))
    pPr.append(pSt)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def make_img_elem(doc, png: bytes, width=IMG_WIDTH):
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.add_run().add_picture(io.BytesIO(png), width=width)
    el = tp._element
    el.getparent().remove(el)
    return el


def insert_after(ref, prose, img, cap):
    ref.addnext(cap)
    ref.addnext(img)
    ref.addnext(prose)


def insert_before(anchor, prose, img, cap):
    parent = anchor.getparent()
    children = list(parent)
    idx = children.index(anchor)
    ref = children[idx - 1] if idx > 0 else None
    if ref is not None:
        ref.addnext(cap)
        ref.addnext(img)
        ref.addnext(prose)
    else:
        parent.insert(0, prose)
        prose.addnext(img)
        img.addnext(cap)


def replace_para_text(p_elem, old: str, new: str) -> bool:
    """Replace old with new across all runs; consolidates text into first w:t."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    t_elems = p_elem.findall(f".//{{{ns}}}t")
    full = "".join(t.text or "" for t in t_elems)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if t_elems:
        t_elems[0].text = new_full
        t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        for t in t_elems[1:]:
            t.text = ""
    return True


# ─────────────────────────────── main ───────────────────────────────


def main():
    sys.stdout.reconfigure(encoding="utf-8")

    # 1. Copy canonical -> v1
    print(f"Copying {CANON.name} -> {V1.name}")
    shutil.copy2(CANON, V1)
    print(f"  Size: {V1.stat().st_size / 1024:.0f} KB")

    # 2. SVG -> PNG
    print("\nConverting SVGs to PNG (300 dpi):")
    png = {
        4: svg_png(FIGS / "FiguraA8_Theta_EN_clean.svg"),
        6: svg_png(FIGS / "Diagram7_Triadic_Tension_clean.svg"),
        7: svg_png(FIGS / "Diagram5_MLOps_Pipeline_clean.svg"),
        8: svg_png(FIGS / "Diagram6_Neurosymbolic_Thermo_clean.svg"),
        9: svg_png(FIGS / "Diagram9_Equity_Map_clean.svg"),
    }

    # 3. Load
    doc = docx.Document(str(V1))
    n0 = len(doc.paragraphs)
    print(f"\nDocument: {n0} paragraphs")

    # 4. Renumber existing Diagram 4->5 and Diagram 5->10
    # IMPORTANT: process Diagram 5->10 FIRST to avoid double-renaming
    print("\nRenumbering existing diagrams...")
    for p in doc.paragraphs:
        if p.style.name == "Caption":
            if p.text.startswith("Diagram 5."):
                if replace_para_text(p._element, "Diagram 5.", "Diagram 10."):
                    print("  Caption: Diagram 5. -> Diagram 10.")
            elif p.text.startswith("Diagram 4."):
                if replace_para_text(p._element, "Diagram 4.", "Diagram 5."):
                    print("  Caption: Diagram 4. -> Diagram 5.")
        # Cross-refs in prose
        if "in Diagram 5 as" in p.text or "in Diagram 5 as" in p.text:
            if replace_para_text(p._element, "in Diagram 5 as", "in Diagram 10 as"):
                print("  Prose xref: Diagram 5 -> Diagram 10")
        if "Diagram 4 illustrates" in p.text:
            if replace_para_text(p._element, "Diagram 4 illustrates", "Diagram 5 illustrates"):
                print("  Prose xref: Diagram 4 -> Diagram 5")

    # 5. Parse prose
    prose = parse_prose(PROSE)
    print(f"\nProse keys loaded: {sorted(prose.keys())}")

    # 6. New diagram captions
    caps = {
        4: (
            "Diagram 4. Instantaneous interference angle θ(t) versus Markovian effective "
            "angle θ_eff across five time steps. The θ(t) series remains stable in "
            "the HITL zone while θ_eff crosses the 120° Circuit Breaker threshold at "
            "t+4, demonstrating the anticipatory governance lead provided by the Markovian "
            "extension (Eq. 5)."
        ),
        6: (
            "Diagram 6. Triadic failure typology: asymmetric partition of normative-violation "
            "space into constitutional failures (absent sovereign predicate), execution-absent-"
            "channel failures (predicate derivable, execution path blocked), and execution-"
            "inertia failures (citation misgrounded or phantom). The geometry encodes the causal "
            "distance from the sovereign predicate axis."
        ),
        7: (
            "Diagram 7. Q-FENG C1 data-flow pipeline (E0–E5): ScopeConfig → "
            "NormChunk → DeonticAtom → ClingoPredicate → SAT/UNSAT verdict → "
            "ψ_S → θ. Pydantic-validated schemas, Parquet persistence, and "
            "deterministic E3/E5 stages ensure full reproducibility. Complements Diagram 1 "
            "(cybernetic governance cycle) at the data-engineering level."
        ),
        8: (
            "Diagram 8. Neurosymbolic free-energy landscape for HITL sovereignty classification. "
            "Sovereign predicates define the constitutional ground-state minimum; elastic "
            "predicates constitute the thermal bath of calibratable regulatory parameters. "
            "The Circuit Breaker threshold corresponds to the energy barrier that the predictor "
            "cannot cross without explicit HITL activation."
        ),
        9: (
            "Diagram 9. Regional equity map of SUS specialist-service distribution across "
            "Brazilian municipalities. The dual annotation layer marks the CF/88 Art. 196 "
            "universal-access gap and the Art. 198 III SUS-regionalisation deficit, quantifying "
            "the constitutional violation that grounds the θ = 134.67° "
            "CIRCUIT_BREAKER outcome of scenario C3 (governance suppression 25.16%)."
        ),
    }

    # 7. Insertions in reverse document order
    insertions = [
        dict(
            num=9, before=False,
            anchor=(
                "This scenario evaluates a structural pattern in Brazilian health policy: "
                "the concentration of SUS specialist services"
            ),
        ),
        dict(
            num=8, before=False,
            anchor=(
                "constitutional failures arise when a required sovereign predicate is absent "
                "from the corpus; execution failures arise when the sovereign predicate exists "
                "but the execution chain is blocked or misgrounded."
            ),
        ),
        dict(
            num=7, before=True,
            anchor=(
                "The pipeline's entry point is the ScopeConfig schema, which parameterises "
                "all subsequent stages"
            ),
        ),
        dict(
            num=6, before=False,
            anchor=(
                "references a non-existent precedent or phantom citation, so the grounding "
                "predicate cannot be derived."
            ),
        ),
        dict(
            num=4, before=False,
            anchor="where γ > 0 is the anticipatory weight and",
        ),
    ]

    for spec in insertions:
        n = spec["num"]
        print(f"\n[INS-{n}] Inserting Diagram {n}...")
        anchor_elem = find_elem(doc, spec["anchor"])
        p_el = make_text_elem(doc, prose.get(f"INS-{n}", f"[prose missing INS-{n}]"), "Normal")
        i_el = make_img_elem(doc, png[n])
        c_el = make_text_elem(doc, caps[n], "Caption")
        if spec["before"]:
            insert_before(anchor_elem, p_el, i_el, c_el)
        else:
            insert_after(anchor_elem, p_el, i_el, c_el)
        print(f"  OK.")

    # 8. Save
    doc.save(str(V1))
    size_mb = V1.stat().st_size / (1024 * 1024)
    nf = len(doc.paragraphs)
    print(f"\nSaved: {V1.name}")
    print(f"Paragraphs: {n0} -> {nf}  (+{nf - n0})")
    print(f"File size: {size_mb:.2f} MB")
    if size_mb > 5.0:
        print("WARNING: file > 5 MB")

    # 9. Post-check
    doc2 = docx.Document(str(V1))
    print("\n=== Post-check: Diagram captions in order ===")
    diagram_order = []
    for i, p in enumerate(doc2.paragraphs):
        if p.style.name == "Caption":
            m = re.match(r"Diagram (\d+)\.", p.text)
            if m:
                dn = int(m.group(1))
                diagram_order.append(dn)
                print(f"  p{i:4d}: Diagram {dn}: {p.text[:70]}")

    print(f"\nDiagram order: {diagram_order}")
    expected = list(range(1, 11))
    if diagram_order == expected:
        print("ORDER CORRECT: 1 through 10")
    else:
        print(f"MISMATCH: expected {expected}")

    print("\n=== Cross-ref audit: Diagram mentions in prose ===")
    for i, p in enumerate(doc2.paragraphs):
        if p.style.name != "Caption":
            for m in re.finditer(r"Diagram \d+", p.text):
                ctx = p.text[max(0, m.start()-30):m.end()+30]
                print(f"  p{i:4d}: ...{ctx}...")

    print("\nDone.")


if __name__ == "__main__":
    main()
