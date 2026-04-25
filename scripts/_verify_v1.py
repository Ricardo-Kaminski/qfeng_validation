"""Verification script for _diagrams_v1.docx final state."""
import re
import docx
from pathlib import Path

V1   = Path(r"C:\Workspace\academico\qfeng_validacao\docs\papers"
            r"\PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed_diagrams_v1.docx")
ORIG = Path(r"C:\Workspace\academico\qfeng_validacao\docs\papers"
            r"\PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")

doc   = docx.Document(str(V1))
orig  = docx.Document(str(ORIG))

print("=== Document statistics ===")
print(f"Original:  {len(orig.paragraphs)} paragraphs  {ORIG.stat().st_size/1024:.0f} KB")
print(f"V1:        {len(doc.paragraphs)} paragraphs  {V1.stat().st_size/1024:.0f} KB")
print(f"Added:     {len(doc.paragraphs) - len(orig.paragraphs)} paragraphs")

print("\n=== Diagram captions (Caption style) ===")
diagram_order = []
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Caption":
        m = re.match(r"(Diagram|Figure|Table) (\d+)\.", p.text)
        if m:
            kind, num = m.group(1), int(m.group(2))
            if kind == "Diagram":
                diagram_order.append(num)
            print(f"  p{i:4d} [{kind} {num:2d}]: {p.text[:70]}")

print(f"\nDiagram order: {diagram_order}")
expected = list(range(1, 11))
print(f"Expected:      {expected}")
print(f"ORDER OK: {diagram_order == expected}")

print("\n=== Cross-reference audit (Diagram N in prose) ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name != "Caption":
        for m in re.finditer(r"Diagram \d+", p.text):
            ctx = p.text[max(0, m.start()-30):m.end()+40]
            print(f"  p{i:4d}: ...{ctx}...")

print("\n=== Image count ===")
from docx.oxml.ns import qn
body = doc.element.body
blips = body.findall(
    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
)
print(f"  Total images embedded: {len(blips)}")

print("\n=== File size ===")
size_mb = V1.stat().st_size / (1024 * 1024)
print(f"  {size_mb:.2f} MB  {'OK' if size_mb <= 5.0 else 'WARNING: > 5 MB'}")
