"""Fix Diagram 4 prose: replace wrong Loss Landscape prose with correct FiguraA8/Markovian prose."""
import docx
from pathlib import Path

V1 = Path(r"C:\Workspace\academico\qfeng_validacao\docs\papers"
          r"\PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed_diagrams_v1.docx")

NEW_PROSE = (
    "Diagram 4 visualises the qualitative distinction between the instantaneous interference "
    "angle θ(t) and the Markovian effective angle θ_eff introduced in Equation 5. "
    "The left portion shows a scenario in which θ(t) remains stable in the intermediate HITL "
    "zone — apparently tolerable under a memoryless monitoring policy. The forward projection "
    "E[θ(t+τ)] traces a deterioration trajectory that, when weighted by the anticipatory "
    "term (γ > 0 in the full form of Eq. A10), drives θ_eff above the 120° "
    "Circuit Breaker threshold. The adaptive sigmoid α(t) — plotted implicitly through "
    "the divergence between the two curves — increases toward 1 as Δpressão(t) > 0, "
    "making the effective angle increasingly sensitive to the deteriorating trajectory. This diagram "
    "provides the geometric rationale for the Markovian extension: a governance framework that "
    "monitors only θ(t) will systematically underestimate normative risk in crisis-onset contexts."
)

WRONG_FRAGMENT = "Diagram 4 illustrates the topology of L_Global as a loss landscape"

doc = docx.Document(str(V1))
fixed = 0
for i, p in enumerate(doc.paragraphs):
    if WRONG_FRAGMENT in p.text and p.style.name != "Caption":
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        t_elems = p._element.findall(f".//{{{ns}}}t")
        if t_elems:
            t_elems[0].text = NEW_PROSE
            t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for t in t_elems[1:]:
                t.text = ""
        print(f"Fixed p{i}: ...{p.text[:80]}...")
        fixed += 1
        break

doc.save(str(V1))
print(f"Fixed {fixed} paragraph(s). Saved.")

# Verify
doc2 = docx.Document(str(V1))
print("\nDiagram 4 mentions in non-caption paragraphs:")
for i, p in enumerate(doc2.paragraphs):
    if "Diagram 4" in p.text and p.style.name != "Caption":
        print(f"  p{i:4d}: {p.text[:120]}")
