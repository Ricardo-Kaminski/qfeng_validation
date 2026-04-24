"""Apply remaining ARS corrections to canonical paper DOCX.

ARS-H4: ψ_N source classification disclosure — add to §3 (after vector definitions)
ARS-H5: Theory of institutional change — add paragraph in §7 before §7.5
ARS-M1: LLM named in E2 — add model name and software table note to Para 164
"""
from __future__ import annotations

import datetime
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("docs/papers/PAPER1_QFENG_FINAL_prob_dados_clingo_auditfixed.docx")


def backup(path: Path) -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dst = path.with_suffix(f".bak_{ts}.docx")
    shutil.copy2(path, dst)
    print(f"Backup: {dst}")
    return dst


def para_idx(doc: docx.Document, fragment: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    raise ValueError(f"Fragment not found: {fragment!r}")


def replace_para_text(doc: docx.Document, idx: int, new_text: str) -> None:
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)
    print(f"  Para {idx} replaced.")


def insert_paragraph_after(doc: docx.Document, ref_idx: int,
                            text: str, style_id: str = "Normal") -> None:
    ref_para = doc.paragraphs[ref_idx]
    new_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.append(pStyle)
    new_para.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    ref_para._element.addnext(new_para)


# ══════════════════════════════════════════════════════════════════════
# ARS-H4: ψ_N source classification in §3
# ══════════════════════════════════════════════════════════════════════

PSI_N_SOURCE_PARA = (
    "\tThe ψ_N vectors in this PoC are constructed by three methods, transparently "
    "disclosed here to ensure reproducibility. (1) Synthetic-calibrated from literature: "
    "C3 ([0.900, 0.070, 0.030] calibrated from 27 corpus normative documents on regional SUS "
    "allocation patterns), C7 ([0.850, 0.100, 0.050] calibrated from the 48,784-record "
    "Obermeyer et al. 2019 dataset), and all CLT scenarios (T-CLT-01 through T-CLT-04) — "
    "ψ_N is calibrated from documented real-world system behaviour but not produced by a "
    "live model inference. (2) Pressure-score interpolated (time series): C2 monthly series "
    "— ψ_N(t) is derived from the composite pressure score (hospital mortality 50%, ICU "
    "utilisation 30%, respiratory disease 20%) at each time step, producing a 12-point "
    "time-varying vector. (3) Predictor-derived (live inference): no scenario in the "
    "current PoC — the C4 Ollama/qwen2.5:14b LLM predictor integration is pending "
    "(Paper 2). The raw pre-normalisation ψ_N values are in "
    "src/qfeng/e5_symbolic/psi_builder.py:_PSI_N_RAW; the time-varying series is in "
    "src/qfeng/e5_symbolic/manaus_sih_loader.py."
)

# ══════════════════════════════════════════════════════════════════════
# ARS-H5: Theory of institutional change — §7 before §7.5
# ══════════════════════════════════════════════════════════════════════

INST_CHANGE_PARA = (
    "A grounding note on the theory of institutional change: Q-FENG is not premised on a "
    "theory that formal normative alignment monitoring is sufficient to produce institutional "
    "change. The Circuit Breaker regime and GSP values reported here are diagnoses, not "
    "remedies. The Manaus collapse of January 2021 was not caused by the absence of a "
    "governance monitoring tool — it was caused by institutional inertia, fragmented "
    "federalism, and resource scarcity. Q-FENG's contribution is epistemic: it formalises "
    "the moment at which the normative architecture becomes calculably opposed to the "
    "predictor's output, generating a falsifiable signal for human interveners. Whether "
    "institutions respond to that signal is a sociological question beyond the scope of "
    "this paper. The framework does not assume — contra rational-choice institutionalism — "
    "that governance actors respond optimally to formal misalignment signals; it assumes "
    "only that such signals are a necessary (though not sufficient) precondition for "
    "evidence-based intervention. The full theoretical grounding — including the VSM "
    "cybernetic rationale for why System 4 prospective intelligence enables earlier "
    "Circuit Breaker activation — is developed in Kaminski (2026a)."
)

# ══════════════════════════════════════════════════════════════════════
# ARS-M1: E2 — add software note to Para 164
# ══════════════════════════════════════════════════════════════════════

E2_LLM_PARA_FRAGMENT = "The extraction used claude-sonnet-4-6 via litellm"

E2_SOFTWARE_ADDENDUM = (
    "\tSoftware stack for the E2 extraction: Python 3.11, litellm 1.x (LLM abstraction "
    "layer), claude-sonnet-4-6 (Anthropic API backend for production; "
    "ollama/qwen2.5:14b for local/zero-cost re-runs), Clingo 5.8.0 (ASP solver), "
    "python-docx 1.1.x (corpus ingestion), pandas 2.x + pyarrow 15.x (parquet "
    "persistence). All random seeds fixed to 1 (Clingo) or 42 (NumPy/bootstrap) "
    "for reproducibility. Full dependency list: pyproject.toml."
)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    backup(DOCX)
    doc = docx.Document(str(DOCX))

    # ── ARS-H4: ψ_N source classification ────────────────────────────
    print("\n[ARS-H4] Adding ψ_N source classification to §3...")
    # Insert after Para 99: "Both vectors are normalised to unit length..."
    idx_psi_n = para_idx(doc, "Both vectors are normalised to unit length")
    print(f"  ψ_N base para: {idx_psi_n}")
    insert_paragraph_after(doc, idx_psi_n, PSI_N_SOURCE_PARA, "Normal")
    print(f"  ψ_N source para inserted after {idx_psi_n}.")

    # ── ARS-H5: Institutional change ─────────────────────────────────
    print("\n[ARS-H5] Adding institutional change paragraph to §7...")
    # Insert before §7.5 heading "7.5 Publication Ecosystem"
    idx_75 = para_idx(doc, "7.5 Publication Ecosystem")
    print(f"  §7.5 heading at para {idx_75}")
    # Insert before §7.5 = after para idx_75 - 1
    insert_paragraph_after(doc, idx_75 - 1, INST_CHANGE_PARA, "Normal")
    print(f"  Institutional change para inserted before §7.5.")

    # ── ARS-M1: E2 software addendum ─────────────────────────────────
    print("\n[ARS-M1] Adding E2 software stack note...")
    idx_e2_llm = para_idx(doc, E2_LLM_PARA_FRAGMENT)
    print(f"  E2 LLM para at {idx_e2_llm}")
    insert_paragraph_after(doc, idx_e2_llm, E2_SOFTWARE_ADDENDUM, "Normal")
    print(f"  Software addendum inserted after {idx_e2_llm}.")

    doc.save(str(DOCX))
    print(f"\nSaved: {DOCX}")

    # Post-check
    doc2 = docx.Document(str(DOCX))
    checks = [
        ("Synthetic-calibrated from literature", "ψ_N source classification"),
        ("Pressure-score interpolated", "ψ_N time-series category"),
        ("Predictor-derived (live inference)", "ψ_N predictor category"),
        ("_PSI_N_RAW", "ψ_N code reference"),
        ("not premised on a theory that formal normative alignment", "institutional change"),
        ("epistemic: it formalises", "institutional change framing"),
        ("Clingo 5.8.0", "software stack disclosure"),
        ("ollama/qwen2.5:14b for local", "LLM named in E2"),
    ]
    all_text = "\n".join(p.text for p in doc2.paragraphs)
    print("\nPost-check:")
    ok = True
    for frag, label in checks:
        found = frag in all_text
        status = "OK" if found else "MISSING"
        print(f"  {status}: {label}")
        if not found:
            ok = False
    assert ok, "Post-check failed."
    print("\nAll checks passed.")


if __name__ == "__main__":
    main()
